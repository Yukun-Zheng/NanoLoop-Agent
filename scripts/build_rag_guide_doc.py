"""Build the NanoLoop Agent RAG and retrieval development guide DOCX."""

from __future__ import annotations

import argparse
import shutil
import subprocess
import tempfile
from datetime import UTC, datetime
from pathlib import Path

import build_v3_handoff_doc as base
from docx import Document
from docx.document import Document as DocumentObject
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.shared import Inches, Pt
from docx.text.paragraph import Paragraph

REPO_ROOT = Path(__file__).resolve().parents[1]
SOURCE_PATH = REPO_ROOT / "docs" / "RAG_RETRIEVAL_DEVELOPMENT_GUIDE.md"
OUTPUT_PATH = REPO_ROOT / "docs" / "NanoLoop_Agent_RAG与检索功能开发指南_v1.0.docx"
CJK_FONT = "Arial Unicode MS"

_PAGE_BREAK_HEADINGS = {
    "1. 阅读方式与今晚必须产出的决策",
}


def _header_paragraph(paragraph: Paragraph) -> None:
    base._clear_paragraph(paragraph)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.tab_stops.add_tab_stop(
        Inches(6.5), alignment=WD_TAB_ALIGNMENT.RIGHT
    )
    left = paragraph.add_run("NanoLoop Agent · RAG 与检索开发指南")
    left.font.bold = True
    left.font.size = Pt(8.5)
    left.font.color.rgb = base._rgb(base.MUTED)
    base._set_fonts(left, cjk=CJK_FONT)
    right = paragraph.add_run("\tv1.0 · 2026-07-18")
    right.font.size = Pt(8.5)
    right.font.color.rgb = base._rgb(base.MUTED)
    base._set_fonts(right, cjk=CJK_FONT)


def _configure_headers_and_footers(doc: DocumentObject) -> None:
    for section in doc.sections:
        _header_paragraph(section.header.paragraphs[0])
        base._footer_paragraph(section.footer.paragraphs[0])
        base._clear_paragraph(section.first_page_header.paragraphs[0])
        base._footer_paragraph(section.first_page_footer.paragraphs[0])


def _style_guide_paragraphs(doc: DocumentObject) -> None:
    base._style_paragraphs(doc)
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if text == "RAG 与检索功能开发指南 v1.0":
            paragraph.style = doc.styles["NanoLoop Cover Subtitle"]
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        elif text == "基于现有 yukun 分支的架构决策、真实资产接入、质量评测与团队任务书":
            paragraph.style = doc.styles["NanoLoop Cover Kicker"]
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        if paragraph.style.name == "Source Code":
            paragraph.paragraph_format.keep_together = True
        if text in _PAGE_BREAK_HEADINGS:
            paragraph.paragraph_format.page_break_before = True


def _apply_cjk_font(doc: DocumentObject) -> None:
    """Use a LibreOffice-renderable CJK font for every document story."""

    for style in doc.styles:
        if style.type == WD_STYLE_TYPE.PARAGRAPH:
            base._set_style_fonts(style, latin=base.LATIN_FONT, cjk=CJK_FONT)

    paragraphs = list(doc.paragraphs)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                paragraphs.extend(cell.paragraphs)
    for section in doc.sections:
        for story in (
            section.header,
            section.footer,
            section.first_page_header,
            section.first_page_footer,
        ):
            paragraphs.extend(story.paragraphs)
    for paragraph in paragraphs:
        for run in paragraph.runs:
            latin = base.CODE_FONT if paragraph.style.name == "Source Code" else base.LATIN_FONT
            base._set_fonts(run, latin=latin, cjk=CJK_FONT)


def _set_metadata(doc: DocumentObject, timestamp: datetime) -> None:
    base._set_metadata(doc, timestamp)
    props = doc.core_properties
    props.title = "NanoLoop Agent RAG 与检索功能开发指南 v1.0"
    props.subject = "基于现有 yukun 分支的真实资产接入、检索评测与团队任务书"
    props.keywords = "NanoLoop Agent, RAG, retrieval, embedding, FAISS, FTS5, handoff"
    props.comments = "Generated from docs/RAG_RETRIEVAL_DEVELOPMENT_GUIDE.md."


def build(source: Path, output: Path, pandoc: str) -> None:
    if not source.is_file():
        raise FileNotFoundError(source)
    timestamp = base._source_timestamp(source).astimezone(UTC)
    output.parent.mkdir(parents=True, exist_ok=True)
    with tempfile.TemporaryDirectory(prefix="nanoloop-rag-guide-") as temp_dir:
        intermediate = Path(temp_dir) / "pandoc.docx"
        subprocess.run(
            [
                pandoc,
                "--from=markdown+pipe_tables+raw_attribute+task_lists+fenced_code_blocks",
                "--to=docx",
                "--standalone",
                f"--resource-path={source.parent}",
                f"--output={intermediate}",
                str(source),
            ],
            check=True,
        )
        doc = Document(intermediate)
        _set_metadata(doc, timestamp)
        base._set_section_geometry(doc)
        base._configure_styles(doc)
        base._patch_numbering(doc)
        _style_guide_paragraphs(doc)
        base._style_tables(doc)
        _configure_headers_and_footers(doc)
        _apply_cjk_font(doc)
        if doc.paragraphs:
            base._remove_trailing_rule(doc.paragraphs[-1])
        doc.save(output)
        base._remove_stale_extended_statistics(output)
        base._remove_empty_comments_part(output, timestamp)


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("--source", type=Path, default=SOURCE_PATH)
    parser.add_argument("--output", type=Path, default=OUTPUT_PATH)
    parser.add_argument("--pandoc", default=shutil.which("pandoc"))
    args = parser.parse_args()
    if not args.pandoc:
        parser.error("pandoc is required to build the RAG guide")
    return args


def main() -> None:
    args = parse_args()
    build(args.source.resolve(), args.output.resolve(), args.pandoc)
    print(args.output.resolve())


if __name__ == "__main__":
    main()
