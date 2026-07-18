"""Build the code-based NanoLoop Agent v3 developer handoff DOCX."""

from __future__ import annotations

import argparse
import shutil
import subprocess
import tempfile
import zipfile
from pathlib import Path
from xml.etree import ElementTree

from docx import Document
from docx.document import Document as DocumentObject
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from docx.styles.style import _ParagraphStyle
from docx.table import Table, _Cell
from docx.text.paragraph import Paragraph
from docx.text.run import Run

REPO_ROOT = Path(__file__).resolve().parents[1]
SOURCE_PATH = REPO_ROOT / "docs" / "NanoLoop_Agent_协同开发规格与接口总文档_v3.0.md"
OUTPUT_PATH = REPO_ROOT / "docs" / "NanoLoop_Agent_协同开发规格与接口总文档_v3.0.docx"

# compact_reference_guide tokens. Arial is the named availability override for the
# preset's Calibri because the build host does not ship Calibri. PingFang SC is the
# named CJK glyph override. Dense table and code sizes are named component overrides.
LATIN_FONT = "Arial"
CJK_FONT = "PingFang SC"
CODE_FONT = "Menlo"
NAVY = "0B2545"
DARK_BLUE = "1F4D78"
HEADING_BLUE = "2E74B5"
MUTED = "596575"
TABLE_HEADER_FILL = "E8EEF5"
CALLOUT_FILL = "F4F6F9"
BORDER = "B8C4D1"
WHITE = "FFFFFF"
CONTENT_WIDTH_DXA = 9360
TABLE_INDENT_DXA = 120


def _rgb(hex_value: str) -> RGBColor:
    return RGBColor.from_string(hex_value)


def _set_fonts(run: Run, *, latin: str = LATIN_FONT, cjk: str = CJK_FONT) -> None:
    run.font.name = latin
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.insert(0, r_fonts)
    r_fonts.set(qn("w:ascii"), latin)
    r_fonts.set(qn("w:hAnsi"), latin)
    r_fonts.set(qn("w:eastAsia"), cjk)
    r_fonts.set(qn("w:cs"), latin)
    lang = r_pr.find(qn("w:lang"))
    if lang is None:
        lang = OxmlElement("w:lang")
        r_pr.append(lang)
    lang.set(qn("w:val"), "zh-CN")
    lang.set(qn("w:eastAsia"), "zh-CN")


def _set_style_fonts(style: _ParagraphStyle, *, latin: str, cjk: str) -> None:
    style.font.name = latin
    r_pr = style.element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.insert(0, r_fonts)
    r_fonts.set(qn("w:ascii"), latin)
    r_fonts.set(qn("w:hAnsi"), latin)
    r_fonts.set(qn("w:eastAsia"), cjk)
    r_fonts.set(qn("w:cs"), latin)
    lang = r_pr.find(qn("w:lang"))
    if lang is None:
        lang = OxmlElement("w:lang")
        r_pr.append(lang)
    lang.set(qn("w:val"), "zh-CN")
    lang.set(qn("w:eastAsia"), "zh-CN")


def _style(
    doc: DocumentObject,
    name: str,
    *,
    size: float,
    color: str = NAVY,
    bold: bool = False,
    before: float = 0,
    after: float = 0,
    line_spacing: float = 1.0,
    keep_with_next: bool = False,
) -> _ParagraphStyle:
    styles = doc.styles
    style = (
        styles[name]
        if name in styles
        else styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
    )
    if not isinstance(style, _ParagraphStyle):
        raise TypeError(f"{name!r} is not a paragraph style")
    _set_style_fonts(style, latin=LATIN_FONT, cjk=CJK_FONT)
    style.font.size = Pt(size)
    style.font.color.rgb = _rgb(color)
    style.font.bold = bold
    fmt = style.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line_spacing
    fmt.keep_with_next = keep_with_next
    fmt.widow_control = True
    return style


def _configure_styles(doc: DocumentObject) -> None:
    normal = _style(doc, "Normal", size=11, color=NAVY, after=6, line_spacing=1.25)
    normal.paragraph_format.keep_with_next = False

    _style(
        doc,
        "Heading 1",
        size=16,
        color=HEADING_BLUE,
        bold=True,
        before=18,
        after=10,
        line_spacing=1.0,
        keep_with_next=True,
    )
    _style(
        doc,
        "Heading 2",
        size=13,
        color=HEADING_BLUE,
        bold=True,
        before=14,
        after=7,
        line_spacing=1.0,
        keep_with_next=True,
    )
    _style(
        doc,
        "Heading 3",
        size=12,
        color=DARK_BLUE,
        bold=True,
        before=10,
        after=5,
        line_spacing=1.0,
        keep_with_next=True,
    )
    _style(
        doc,
        "NanoLoop Cover Title",
        size=30,
        color=NAVY,
        bold=True,
        before=6,
        after=5,
        line_spacing=1.0,
        keep_with_next=True,
    )
    _style(
        doc,
        "NanoLoop Cover Subtitle",
        size=16,
        color=DARK_BLUE,
        bold=False,
        before=0,
        after=5,
        line_spacing=1.05,
        keep_with_next=True,
    )
    _style(
        doc,
        "NanoLoop Cover Kicker",
        size=10.5,
        color=MUTED,
        bold=True,
        before=0,
        after=10,
        line_spacing=1.0,
        keep_with_next=True,
    )
    _style(
        doc,
        "NanoLoop Cover Label",
        size=11.5,
        color=DARK_BLUE,
        bold=True,
        before=8,
        after=5,
        line_spacing=1.0,
        keep_with_next=True,
    )
    _style(
        doc,
        "NanoLoop Table",
        size=8.6,
        color=NAVY,
        after=0,
        line_spacing=1.05,
    )
    _style(
        doc,
        "NanoLoop Table Header",
        size=8.7,
        color=NAVY,
        bold=True,
        after=0,
        line_spacing=1.0,
    )
    _style(
        doc,
        "NanoLoop Code Block",
        size=8.2,
        color=NAVY,
        after=3,
        line_spacing=1.0,
    )

    if "Source Code" in doc.styles:
        source = doc.styles["Source Code"]
        if isinstance(source, _ParagraphStyle):
            _set_style_fonts(source, latin=CODE_FONT, cjk=CJK_FONT)
            source.font.size = Pt(8.2)
            source.font.color.rgb = _rgb(NAVY)
            source.paragraph_format.space_before = Pt(3)
            source.paragraph_format.space_after = Pt(3)
            source.paragraph_format.line_spacing = 1.0

    if "Verbatim Char" in doc.styles:
        verbatim = doc.styles["Verbatim Char"]
        verbatim.font.name = CODE_FONT
        verbatim.font.size = Pt(8.4)
        verbatim.font.color.rgb = _rgb(DARK_BLUE)
        r_pr = verbatim.element.get_or_add_rPr()
        r_fonts = r_pr.rFonts
        if r_fonts is None:
            r_fonts = OxmlElement("w:rFonts")
            r_pr.insert(0, r_fonts)
        for key in ("ascii", "hAnsi", "cs"):
            r_fonts.set(qn(f"w:{key}"), CODE_FONT)
        r_fonts.set(qn("w:eastAsia"), CJK_FONT)

    if "List Paragraph" in doc.styles:
        list_style = doc.styles["List Paragraph"]
        if isinstance(list_style, _ParagraphStyle):
            _set_style_fonts(list_style, latin=LATIN_FONT, cjk=CJK_FONT)
            list_style.font.size = Pt(11)
            list_style.font.color.rgb = _rgb(NAVY)
            fmt = list_style.paragraph_format
            fmt.left_indent = Inches(0.375)
            fmt.first_line_indent = Inches(-0.188)
            fmt.space_after = Pt(4)
            fmt.line_spacing = 1.25

    if "Quote" in doc.styles:
        quote = doc.styles["Quote"]
        if isinstance(quote, _ParagraphStyle):
            _set_style_fonts(quote, latin=LATIN_FONT, cjk=CJK_FONT)
            quote.font.size = Pt(10.2)
            quote.font.color.rgb = _rgb(DARK_BLUE)
            quote.paragraph_format.left_indent = Inches(0.18)
            quote.paragraph_format.right_indent = Inches(0.05)
            quote.paragraph_format.space_before = Pt(4)
            quote.paragraph_format.space_after = Pt(8)
            quote.paragraph_format.line_spacing = 1.15


def _set_section_geometry(doc: DocumentObject) -> None:
    for section in doc.sections:
        section.page_width = Inches(8.5)
        section.page_height = Inches(11)
        section.top_margin = Inches(1.0)
        section.right_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.header_distance = Inches(0.492)
        section.footer_distance = Inches(0.492)
        section.different_first_page_header_footer = True


def _clear_paragraph(paragraph: Paragraph) -> None:
    for child in list(paragraph._p):
        paragraph._p.remove(child)


def _page_field(paragraph: Paragraph) -> None:
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    result = OxmlElement("w:t")
    result.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instruction, separate, result, end])
    _set_fonts(run)
    run.font.size = Pt(8.5)
    run.font.color.rgb = _rgb(MUTED)


def _header_paragraph(paragraph: Paragraph) -> None:
    _clear_paragraph(paragraph)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.tab_stops.add_tab_stop(
        Inches(6.5), alignment=WD_TAB_ALIGNMENT.RIGHT
    )
    left = paragraph.add_run("NanoLoop Agent · 协同开发手册")
    left.font.bold = True
    left.font.size = Pt(8.5)
    left.font.color.rgb = _rgb(MUTED)
    _set_fonts(left)
    right = paragraph.add_run("\tv3.0 · 2026-07-18")
    right.font.size = Pt(8.5)
    right.font.color.rgb = _rgb(MUTED)
    _set_fonts(right)


def _footer_paragraph(paragraph: Paragraph) -> None:
    _clear_paragraph(paragraph)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    prefix = paragraph.add_run("NanoLoop Agent  ·  Page ")
    prefix.font.size = Pt(8.5)
    prefix.font.color.rgb = _rgb(MUTED)
    _set_fonts(prefix)
    _page_field(paragraph)


def _configure_headers_and_footers(doc: DocumentObject) -> None:
    for section in doc.sections:
        _header_paragraph(section.header.paragraphs[0])
        _footer_paragraph(section.footer.paragraphs[0])
        first_header = section.first_page_header.paragraphs[0]
        _clear_paragraph(first_header)
        _footer_paragraph(section.first_page_footer.paragraphs[0])


def _set_cell_margins(cell: _Cell) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for edge, value in (("top", 80), ("start", 120), ("bottom", 80), ("end", 120)):
        node = tc_mar.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def _shade_cell(cell: _Cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = tc_pr.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        tc_pr.append(shading)
    shading.set(qn("w:val"), "clear")
    shading.set(qn("w:color"), "auto")
    shading.set(qn("w:fill"), fill)


def _set_table_borders(table: Table) -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        node = borders.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            borders.append(node)
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), "4")
        node.set(qn("w:space"), "0")
        node.set(qn("w:color"), BORDER)


def _column_widths(table: Table) -> list[int]:
    count = len(table.columns)
    if count == 0 and table.rows:
        count = max(len(row.cells) for row in table.rows)
    if count == 0:
        return []
    headers = [cell.text.strip() for cell in table.rows[0].cells] if table.rows else []
    first = headers[0] if headers else ""
    if count == 2:
        return [2000, 7360]
    if count == 3:
        if first in {"版本", "层级"}:
            return [1450, 2850, 5060]
        return [1900, 3150, 4310]
    if count == 4:
        if first == "方法":
            return [800, 2950, 4050, 1560]
        if first == "变量":
            return [2450, 1550, 900, 4460]
        if first in {"角色", "开发者"}:
            return [1050, 2050, 3100, 3160]
        if first == "范畴":
            return [1550, 1450, 3250, 3110]
        if first == "优先级":
            return [1050, 2850, 5460]
        return [1400, 2250, 2900, 2810]
    base = CONTENT_WIDTH_DXA // max(count, 1)
    widths = [base] * count
    widths[-1] += CONTENT_WIDTH_DXA - sum(widths)
    return widths


def _set_table_geometry(table: Table, widths: list[int]) -> None:
    column_count = len(table.columns)
    if column_count == 0 and table.rows:
        column_count = max(len(row.cells) for row in table.rows)
    if len(widths) != column_count or sum(widths) != CONTENT_WIDTH_DXA:
        raise ValueError("table widths must match columns and sum to 9360 DXA")
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr

    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(CONTENT_WIDTH_DXA))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(TABLE_INDENT_DXA))
    tbl_ind.set(qn("w:type"), "dxa")

    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")

    grid = table._tbl.tblGrid
    grid_columns = list(grid.iterchildren(tag=qn("w:gridCol")))
    while len(grid_columns) < len(widths):
        node = OxmlElement("w:gridCol")
        grid.append(node)
        grid_columns.append(node)
    for node, width in zip(grid_columns, widths, strict=True):
        node.set(qn("w:w"), str(width))

    for row in table.rows:
        for index, cell in enumerate(row.cells):
            width = widths[min(index, len(widths) - 1)]
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")


def _repeat_header(row: object) -> None:
    tr = row._tr  # type: ignore[attr-defined]
    tr_pr = tr.get_or_add_trPr()
    header = tr_pr.find(qn("w:tblHeader"))
    if header is None:
        header = OxmlElement("w:tblHeader")
        tr_pr.append(header)
    header.set(qn("w:val"), "true")


def _prevent_row_split(row: object) -> None:
    tr = row._tr  # type: ignore[attr-defined]
    tr_pr = tr.get_or_add_trPr()
    cant_split = tr_pr.find(qn("w:cantSplit"))
    if cant_split is None:
        cant_split = OxmlElement("w:cantSplit")
        tr_pr.append(cant_split)
    cant_split.set(qn("w:val"), "true")


def _style_tables(doc: DocumentObject) -> None:
    for table_index, table in enumerate(doc.tables):
        widths = _column_widths(table)
        if not widths:
            continue
        _set_table_geometry(table, widths)
        _set_table_borders(table)
        if table.rows:
            _repeat_header(table.rows[0])
        cover_override = table_index < 2
        for row_index, row in enumerate(table.rows):
            _prevent_row_split(row)
            for cell in row.cells:
                cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                _set_cell_margins(cell)
                if row_index == 0:
                    _shade_cell(cell, DARK_BLUE if cover_override else TABLE_HEADER_FILL)
                for paragraph in cell.paragraphs:
                    paragraph.style = (
                        doc.styles["NanoLoop Table Header"]
                        if row_index == 0
                        else doc.styles["NanoLoop Table"]
                    )
                    paragraph.paragraph_format.keep_together = False
                    paragraph.paragraph_format.keep_with_next = False
                    for run in paragraph.runs:
                        _set_fonts(run)
                        run.font.size = Pt(8.7 if row_index == 0 else 8.6)
                        if row_index == 0:
                            run.font.bold = True
                            run.font.color.rgb = _rgb(WHITE if cover_override else NAVY)


def _patch_numbering(doc: DocumentObject) -> None:
    numbering = doc.part.numbering_part.element
    for level in numbering.iter(qn("w:lvl")):
        ilvl = int(level.get(qn("w:ilvl"), "0"))
        left = 540 + ilvl * 360
        hanging = 270
        p_pr = level.find(qn("w:pPr"))
        if p_pr is None:
            p_pr = OxmlElement("w:pPr")
            level.append(p_pr)
        tabs = p_pr.find(qn("w:tabs"))
        if tabs is None:
            tabs = OxmlElement("w:tabs")
            p_pr.append(tabs)
        for old in list(tabs):
            tabs.remove(old)
        tab = OxmlElement("w:tab")
        tab.set(qn("w:val"), "num")
        tab.set(qn("w:pos"), str(left))
        tabs.append(tab)
        indent = p_pr.find(qn("w:ind"))
        if indent is None:
            indent = OxmlElement("w:ind")
            p_pr.append(indent)
        indent.set(qn("w:left"), str(left))
        indent.set(qn("w:hanging"), str(hanging))
        spacing = p_pr.find(qn("w:spacing"))
        if spacing is None:
            spacing = OxmlElement("w:spacing")
            p_pr.append(spacing)
        spacing.set(qn("w:after"), "80")
        spacing.set(qn("w:line"), "300")
        spacing.set(qn("w:lineRule"), "auto")
        num_fmt = level.find(qn("w:numFmt"))
        if num_fmt is not None and num_fmt.get(qn("w:val")) == "bullet":
            level_text = level.find(qn("w:lvlText"))
            if level_text is not None:
                level_text.set(qn("w:val"), "•")


def _style_quote_paragraph(paragraph: Paragraph) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    shading = p_pr.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        p_pr.append(shading)
    shading.set(qn("w:val"), "clear")
    shading.set(qn("w:fill"), CALLOUT_FILL)
    borders = p_pr.find(qn("w:pBdr"))
    if borders is None:
        borders = OxmlElement("w:pBdr")
        p_pr.append(borders)
    left = borders.find(qn("w:left"))
    if left is None:
        left = OxmlElement("w:left")
        borders.append(left)
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), "14")
    left.set(qn("w:space"), "7")
    left.set(qn("w:color"), HEADING_BLUE)


def _style_paragraphs(doc: DocumentObject) -> None:
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if text == "NanoLoop Agent":
            paragraph.style = doc.styles["NanoLoop Cover Title"]
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        elif text == "协同开发规格与接口总文档 v3.0":
            paragraph.style = doc.styles["NanoLoop Cover Subtitle"]
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        elif text == "面向现有仓库的开发者接手、修改与扩展手册":
            paragraph.style = doc.styles["NanoLoop Cover Kicker"]
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        elif text == "本版交付信号":
            paragraph.style = doc.styles["NanoLoop Cover Label"]
        elif paragraph.style.name == "Quote":
            _style_quote_paragraph(paragraph)
        elif paragraph.style.name == "Source Code":
            for run in paragraph.runs:
                _set_fonts(run, latin=CODE_FONT, cjk=CJK_FONT)
                run.font.size = Pt(8.2)
        else:
            for run in paragraph.runs:
                if run.style is None or run.style.name != "Verbatim Char":
                    _set_fonts(run)

        if paragraph.style.name.startswith("Heading"):
            paragraph.paragraph_format.keep_with_next = True
            paragraph.paragraph_format.keep_together = True
        paragraph.paragraph_format.widow_control = True


def _set_metadata(doc: DocumentObject) -> None:
    props = doc.core_properties
    props.title = "NanoLoop Agent 协同开发规格与接口总文档 v3.0"
    props.subject = "基于当前仓库代码的开发者接手、修改与扩展手册"
    props.author = "NanoLoop Agent Team"
    props.last_modified_by = "NanoLoop Agent Team"
    props.keywords = "NanoLoop Agent, SEM, developer handoff, API, RAG, inference"
    props.comments = "Generated from the repository Markdown source."
    settings = doc.settings.element
    update = settings.find(qn("w:updateFields"))
    if update is None:
        update = OxmlElement("w:updateFields")
        settings.append(update)
    update.set(qn("w:val"), "true")


def _remove_trailing_rule(paragraph: Paragraph) -> None:
    if paragraph.text.strip():
        return
    p_pr = paragraph._p.pPr
    if p_pr is None:
        return
    borders = p_pr.find(qn("w:pBdr"))
    if borders is not None:
        p_pr.remove(borders)


def _remove_stale_extended_statistics(path: Path) -> None:
    """Remove template-derived pagination statistics that Word may not refresh."""

    app_part = "docProps/app.xml"
    stats = {
        "Characters",
        "CharactersWithSpaces",
        "Lines",
        "Pages",
        "Paragraphs",
        "TotalTime",
        "Words",
    }
    with zipfile.ZipFile(path, "r") as source_zip:
        entries = [(item, source_zip.read(item.filename)) for item in source_zip.infolist()]

    extended_ns = (
        "http://schemas.openxmlformats.org/officeDocument/2006/extended-properties"
    )
    ElementTree.register_namespace("", extended_ns)
    temp_path = path.with_suffix(f"{path.suffix}.tmp")
    try:
        with zipfile.ZipFile(temp_path, "w") as target_zip:
            for item, payload in entries:
                if item.filename == app_part:
                    root = ElementTree.fromstring(payload)
                    for child in list(root):
                        if child.tag.removeprefix(f"{{{extended_ns}}}") in stats:
                            root.remove(child)
                    payload = ElementTree.tostring(
                        root,
                        encoding="utf-8",
                        xml_declaration=True,
                    )
                target_zip.writestr(item, payload)
        temp_path.replace(path)
    finally:
        temp_path.unlink(missing_ok=True)


def build(source: Path, output: Path, pandoc: str) -> None:
    if not source.is_file():
        raise FileNotFoundError(source)
    output.parent.mkdir(parents=True, exist_ok=True)
    with tempfile.TemporaryDirectory(prefix="nanoloop-v3-doc-") as temp_dir:
        intermediate = Path(temp_dir) / "pandoc.docx"
        subprocess.run(
            [
                pandoc,
                "--from=markdown+pipe_tables+raw_attribute+task_lists+fenced_code_blocks",
                "--to=docx",
                "--standalone",
                f"--resource-path={source.parent}",
                f"--output={intermediate}",
                str(source),
            ],
            check=True,
        )
        doc = Document(intermediate)
        _set_metadata(doc)
        _set_section_geometry(doc)
        _configure_styles(doc)
        _patch_numbering(doc)
        _style_paragraphs(doc)
        _style_tables(doc)
        _configure_headers_and_footers(doc)
        if doc.paragraphs:
            _remove_trailing_rule(doc.paragraphs[-1])
        doc.save(output)
        _remove_stale_extended_statistics(output)


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("--source", type=Path, default=SOURCE_PATH)
    parser.add_argument("--output", type=Path, default=OUTPUT_PATH)
    parser.add_argument("--pandoc", default=shutil.which("pandoc"))
    args = parser.parse_args()
    if not args.pandoc:
        parser.error("pandoc is required; install it and rerun make handoff-doc")
    return args


def main() -> None:
    args = parse_args()
    build(args.source.resolve(), args.output.resolve(), args.pandoc)
    print(args.output.resolve())


if __name__ == "__main__":
    main()
