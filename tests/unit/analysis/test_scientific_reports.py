import io
from datetime import UTC, datetime
from pathlib import Path

from docx import Document
from PIL import Image, ImageDraw

from app.agent.unified_query import DataQuery, DataQueryResult
from app.analysis.reporting import JobExportSnapshot
from app.analysis.scientific_reports import ScientificReportBuilder, _batch_summary
from app.contracts.analyses import (
    AnalysisJobDTO,
    AnalysisROI,
    ImageAssetDTO,
    ImageSummaryDTO,
    InferenceOptions,
    PixelRect,
    QualityReportDTO,
    RunConfiguration,
    SegmentationRunDTO,
)
from app.contracts.common import HealthComponent
from app.contracts.enums import JobStatus, QualityStatus, RoiMode
from app.contracts.queries import ToolEvidence
from app.rag.providers import ConversationProviderAnswer
from app.storage.file_store import LocalFileStore
from app.storage.paths import StoragePaths


class RecordingDataTools:
    def __init__(self) -> None:
        self.questions: list[str] = []

    def answer(self, query: DataQuery) -> DataQueryResult:
        self.questions.append(query.question)
        return DataQueryResult(
            answer="已返回当前运行的权威统计。",
            evidence=(
                ToolEvidence(
                    tool_name="run_overview",
                    validated_arguments={"run_ids": list(query.run_ids)},
                    aggregates={
                        "particle_count": 95,
                        "mean_equivalent_diameter_nm": 59.372,
                    },
                    units={
                        "particle_count": "count",
                        "mean_equivalent_diameter_nm": "nm",
                    },
                    source_run_ids=list(query.run_ids),
                    quality_warnings=["small_fragment_ratio_high"],
                ),
            ),
            confidence="high",
            limitations=("单视野结果不能代表完整样品。",),
        )


class GroundedLocalProvider:
    model = "fixture-qwen"

    def health(self) -> HealthComponent:
        return HealthComponent(status="healthy")

    def generate_conversation(self, **_: object) -> ConversationProviderAnswer:
        return ConversationProviderAnswer(
            answer=(
                "自动质量门控不替代人工边界抽查，单个图像视野不能代表完整样品，"
                "建议增加重复视野 [D1]。"
            ),
            used_data_ids=("D1",),
            used_citation_ids=(),
            confidence="high",
        )


def test_scientific_report_builds_preview_docx_and_pdf_from_one_snapshot(
    tmp_path: Path,
) -> None:
    now = datetime.now(UTC)
    analysis_roi = AnalysisROI(valid_rect=PixelRect(x1=0, y1=0, x2=2048, y2=1406))
    inference = InferenceOptions(threshold=0.5, min_area_px=8)
    configuration = RunConfiguration(
        model_id="unet-agglomerated-specialist",
        model_version="1.3.0",
        roi_mode=RoiMode.FULL_IMAGE,
        analysis_roi=analysis_roi,
        inference=inference,
        preprocess_profile="sem_gray_v1",
        postprocess_profile="agglomerated_v1",
        image_sha256="a" * 64,
        scale_nm_per_pixel=1.0,
        created_at=now,
    )
    summary = ImageSummaryDTO(
        run_id="run_report",
        particle_count=95,
        roi_area_px=2_879_488,
        number_density_px2=0.000033,
        number_density_um2=0.033,
        mean_equivalent_diameter_px=59.372,
        mean_equivalent_diameter_nm=59.372,
        coverage_ratio=0.1069,
        perimeter_density_px=0.009164,
        perimeter_density_um=9.164,
        quality_status=QualityStatus.WARN,
    )
    quality = QualityReportDTO(
        status=QualityStatus.WARN,
        reasons=["small_fragment_ratio_high"],
        recommendations=["提高 min_area_px 或复核小颗粒专用模型"],
    )
    run = SegmentationRunDTO(
        run_id="run_report",
        job_id="job_report",
        image_id="image_report",
        model_id=configuration.model_id,
        status=JobStatus.COMPLETED_WITH_WARNINGS,
        roi_mode=RoiMode.FULL_IMAGE,
        threshold=0.5,
        inference=inference,
        configuration=configuration,
        summary=summary,
        quality=quality,
        runtime_ms=1200,
        created_at=now,
        updated_at=now,
    )
    file_store = LocalFileStore(
        StoragePaths(tmp_path / "outputs"),
        max_upload_bytes=10_000_000,
        token_secret=b"x" * 32,
    )
    overlay_path = file_store.paths.run_artifact(
        "job_report",
        "image_report",
        "run_report",
        "overlay.png",
    )
    overlay = Image.new("RGB", (1200, 640), "#30343b")
    drawing = ImageDraw.Draw(overlay)
    for x, y, radius in ((180, 180, 90), (420, 250, 120), (760, 180, 105), (990, 360, 130)):
        drawing.ellipse(
            (x - radius, y - radius, x + radius, y + radius),
            fill="#8a63d2",
            outline="#d9c9ff",
            width=8,
        )
    overlay_bytes = io.BytesIO()
    overlay.save(overlay_bytes, format="PNG")
    file_store.atomic_write_bytes(overlay_path, overlay_bytes.getvalue())

    snapshot = JobExportSnapshot(
        job=AnalysisJobDTO(
            job_id="job_report",
            name="BaNi-3",
            status=JobStatus.COMPLETED_WITH_WARNINGS,
            created_at=now,
            updated_at=now,
        ),
        images=(
            ImageAssetDTO(
                image_id="image_report",
                job_id="job_report",
                filename="BaNi-3.tif",
                sha256="a" * 64,
                width=2048,
                height=1406,
                bit_depth=8,
                sample_id="BaNi-3",
                scale_nm_per_pixel=1.0,
                analysis_roi=analysis_roi,
            ),
        ),
        runs=(run,),
        run_artifact_paths=(("run_report", file_store.paths.relative_path(overlay_path)),),
    )
    data_tools = RecordingDataTools()

    preview, docx_file, pdf_file = ScientificReportBuilder(
        file_store=file_store,
        data_tools=data_tools,
        llm_provider=None,
    ).build(
        snapshot=snapshot,
        tenant_id="tenant_report",
        run_ids=["run_report"],
    )

    assert len(data_tools.questions) == 4
    assert preview.scale_status == "physical"
    assert preview.fallback_used is True
    assert preview.synthesis_provider == "deterministic_fallback"
    assert preview.headline_metrics[2].display_value == "59.37 nm"
    assert preview.quality_status is QualityStatus.WARN
    assert any("min_area_px=8" in item.action for item in preview.recommendations)
    assert docx_file.path.read_bytes().startswith(b"PK")
    assert pdf_file.path.read_bytes().startswith(b"%PDF")

    document = Document(io.BytesIO(docx_file.path.read_bytes()))
    report_text = "\n".join(paragraph.text for paragraph in document.paragraphs)
    assert "BaNi-3 - SEM 纳米颗粒分析报告" in report_text
    assert "限制、不确定性与稳健性" in report_text
    assert "59.37 nm" in "\n".join(
        cell.text for table in document.tables for row in table.rows for cell in row.cells
    )

    second_summary = summary.model_copy(
        update={
            "run_id": "run_report_2",
            "particle_count": 120,
            "coverage_ratio": 0.145,
            "mean_equivalent_diameter_px": 66.0,
            "mean_equivalent_diameter_nm": 66.0,
        }
    )
    second_run = run.model_copy(
        update={
            "run_id": "run_report_2",
            "image_id": "image_report_2",
            "summary": second_summary,
        }
    )
    second_image = snapshot.images[0].model_copy(
        update={
            "image_id": "image_report_2",
            "filename": "BaNi-4.tif",
            "sample_id": "BaNi-4",
            "sha256": "b" * 64,
        }
    )
    batch_snapshot = JobExportSnapshot(
        job=snapshot.job,
        images=(*snapshot.images, second_image),
        runs=(run, second_run),
        run_artifact_paths=snapshot.run_artifact_paths,
    )
    batch_preview, batch_docx, batch_pdf = ScientificReportBuilder(
        file_store=file_store,
        data_tools=data_tools,
        llm_provider=None,
    ).build(
        snapshot=batch_snapshot,
        tenant_id="tenant_report",
        run_ids=["run_report", "run_report_2"],
    )
    assert batch_preview.analysis_mode == "batch"
    assert batch_preview.batch_summary is not None
    assert batch_preview.batch_summary.image_count == 2
    assert batch_preview.batch_summary.run_count == 2
    assert batch_preview.run_summaries[1].filename == "BaNi-4.tif"
    assert "批量分析报告" in batch_preview.title
    assert batch_docx.path.read_bytes().startswith(b"PK")
    assert batch_pdf.path.read_bytes().startswith(b"%PDF")

    llm_summary, synthesis_provider, synthesis_model, fallback_used = ScientificReportBuilder(
        file_store=file_store,
        data_tools=data_tools,
        llm_provider=GroundedLocalProvider(),  # type: ignore[arg-type]
    )._technical_summary(
        snapshot=snapshot,
        runs=(run,),
        evidence=(
            ToolEvidence(
                tool_name="run_overview",
                validated_arguments={"run_ids": ["run_report"]},
                aggregates={"particle_count": 95},
                units={"particle_count": "count"},
                source_run_ids=["run_report"],
            ),
        ),
    )
    assert synthesis_provider == "local_llm"
    assert synthesis_model == "fixture-qwen"
    assert fallback_used is False
    assert "本地模型综合：" in llm_summary
    assert "建议增加重复视野 [D1]" in llm_summary


def test_batch_summary_exposes_distributions_variability_and_iqr_outliers() -> None:
    now = datetime.now(UTC)
    analysis_roi = AnalysisROI(valid_rect=PixelRect(x1=0, y1=0, x2=100, y2=100))
    inference = InferenceOptions(threshold=0.5, min_area_px=8)
    configuration = RunConfiguration(
        model_id="unet-batch",
        model_version="1.0.0",
        roi_mode=RoiMode.FULL_IMAGE,
        analysis_roi=analysis_roi,
        inference=inference,
        preprocess_profile="sem_gray_v1",
        postprocess_profile="default_v1",
        scale_nm_per_pixel=1.0,
        created_at=now,
    )
    runs: list[SegmentationRunDTO] = []
    for index, (count, coverage, diameter) in enumerate(
        (
            (9, 0.09, 18.0),
            (10, 0.10, 19.0),
            (10, 0.10, 20.0),
            (11, 0.11, 21.0),
            (100, 0.90, 90.0),
        ),
        start=1,
    ):
        run_id = f"run_batch_{index}"
        summary = ImageSummaryDTO(
            run_id=run_id,
            particle_count=count,
            roi_area_px=10_000,
            number_density_px2=count / 10_000,
            number_density_um2=count / 0.01,
            mean_equivalent_diameter_px=diameter,
            mean_equivalent_diameter_nm=diameter,
            coverage_ratio=coverage,
            perimeter_density_px=coverage / 10,
            perimeter_density_um=coverage * 100,
            quality_status=QualityStatus.PASS,
        )
        runs.append(
            SegmentationRunDTO(
                run_id=run_id,
                job_id="job_batch",
                image_id=f"image_batch_{index}",
                model_id=configuration.model_id,
                status=JobStatus.COMPLETED,
                roi_mode=RoiMode.FULL_IMAGE,
                inference=inference,
                configuration=configuration,
                summary=summary,
                quality=QualityReportDTO(status=QualityStatus.PASS),
                created_at=now,
                updated_at=now,
            )
        )

    batch = _batch_summary(runs)

    assert batch.image_count == 5
    assert batch.run_count == 5
    assert batch.total_particle_count == 140
    counts = next(item for item in batch.distributions if item.key == "particle_count")
    assert counts.sample_count == 5
    assert counts.median == 10
    assert counts.q1 == 10
    assert counts.q3 == 11
    assert counts.std_dev > 0
    assert counts.coefficient_of_variation is not None
    assert any(
        item.run_id == "run_batch_5"
        and item.metric_key == "particle_count"
        and item.direction == "high"
        for item in batch.outliers
    )
