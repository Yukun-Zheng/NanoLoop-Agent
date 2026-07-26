"""Evidence-first scientific reports with DOCX and PDF renderers.

Numerical statements are assembled from persisted run summaries and the
deterministic data-tool boundary.  The optional local LLM is allowed to rewrite
only the technical summary; its visible evidence markers are validated before
the text is accepted.
"""

from __future__ import annotations

import hashlib
import io
import json
import logging
import re
import threading
from collections.abc import Iterable, Sequence
from datetime import datetime
from html import escape
from pathlib import Path
from statistics import fmean, pstdev
from typing import Any

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER, TA_LEFT
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.cidfonts import UnicodeCIDFont
from reportlab.platypus import (
    Image,
    KeepTogether,
    Paragraph,
    SimpleDocTemplate,
    Spacer,
    Table,
    TableStyle,
)

from app.agent.evidence_validation import validate_conversation_answer
from app.agent.unified_query import DataQuery, DataToolService
from app.analysis.reporting import JobExportSnapshot
from app.contracts.analyses import SegmentationRunDTO
from app.contracts.common import utc_now
from app.contracts.enums import QualityStatus
from app.contracts.queries import Citation, ToolEvidence
from app.contracts.reports import (
    BatchMetricDistributionDTO,
    BatchOutlierDTO,
    BatchResultSummaryDTO,
    ReportFindingDTO,
    ReportMetricDTO,
    ReportProvenanceDTO,
    ReportRecommendationDTO,
    ReportRunSummaryDTO,
    ScientificReportPreviewDTO,
)
from app.rag.providers import (
    AnswerProviderError,
    CitationValidationError,
    OpenAICompatibleProvider,
)
from app.storage.file_store import LocalFileStore, StoredFile

_PDF_FONT = "STSong-Light"
_PDF_FONT_LOCK = threading.Lock()
_REPORT_LOCK = threading.Lock()
_BLUE = RGBColor(46, 116, 181)
_DARK_BLUE = RGBColor(31, 77, 120)
_INK = RGBColor(29, 35, 48)
_MUTED = RGBColor(92, 100, 116)
_TABLE_HEADER_FILL = "F2F4F7"
_CALLOUT_FILL = "F4F6F9"
_DOCX_FONT = "Arial Unicode MS"
logger = logging.getLogger(__name__)


class ScientificReportBuilder:
    """Build one structured preview and the matching immutable document files."""

    def __init__(
        self,
        *,
        file_store: LocalFileStore,
        data_tools: DataToolService,
        llm_provider: OpenAICompatibleProvider | None,
    ) -> None:
        self.file_store = file_store
        self.data_tools = data_tools
        self.llm_provider = llm_provider

    def build(
        self,
        *,
        snapshot: JobExportSnapshot,
        tenant_id: str,
        run_ids: Sequence[str],
    ) -> tuple[ScientificReportPreviewDTO, StoredFile, StoredFile]:
        runs_by_id = {run.run_id: run for run in snapshot.runs}
        selected = tuple(runs_by_id[run_id] for run_id in run_ids)
        if not selected:
            raise ValueError("scientific report requires at least one run")
        if any(run.summary is None or run.quality is None for run in selected):
            raise ValueError("scientific report requires completed run summaries")

        generated_at = utc_now()
        tool_evidence, tool_limitations = self._collect_evidence(
            job_id=snapshot.job.job_id,
            tenant_id=tenant_id,
            run_ids=tuple(run_ids),
        )
        evidence = (
            _report_interpretation_evidence(selected),
            _report_synthesis_evidence(selected),
            *tool_evidence,
        )
        technical_summary, provider, model, fallback_used = self._technical_summary(
            snapshot=snapshot,
            runs=selected,
            evidence=evidence,
        )
        report_id = _report_id(
            snapshot=snapshot,
            run_ids=run_ids,
            generated_at=generated_at,
            technical_summary=technical_summary,
            evidence=evidence,
        )
        preview = _build_preview(
            report_id=report_id,
            snapshot=snapshot,
            runs=selected,
            generated_at=generated_at,
            technical_summary=technical_summary,
            synthesis_provider=provider,
            synthesis_model=model,
            fallback_used=fallback_used,
            evidence=evidence,
            tool_limitations=tool_limitations,
        )
        overlay = _primary_overlay(snapshot, selected[0].run_id, self.file_store)
        with _REPORT_LOCK:
            docx_bytes = _render_docx(preview, overlay=overlay)
            pdf_bytes = _render_pdf(preview, overlay=overlay)
            docx_file = self._store_report(
                snapshot.job.job_id,
                f"nanoloop-scientific-report-{report_id[-12:]}.docx",
                docx_bytes,
            )
            pdf_file = self._store_report(
                snapshot.job.job_id,
                f"nanoloop-scientific-report-{report_id[-12:]}.pdf",
                pdf_bytes,
            )
        return preview, docx_file, pdf_file

    def _collect_evidence(
        self,
        *,
        job_id: str,
        tenant_id: str,
        run_ids: tuple[str, ...],
    ) -> tuple[tuple[ToolEvidence, ...], tuple[str, ...]]:
        questions = (
            "概括当前任务的颗粒数量、数量密度、平均粒径、覆盖率和周长密度",
            "分析平均粒径分布、中位数和四分位",
            "找出当前运行中的异常和质量限制",
            "哪些运行需要复核并说明原因",
        )
        evidence: list[ToolEvidence] = []
        limitations: list[str] = []
        seen: set[str] = set()
        for question in questions:
            result = self.data_tools.answer(
                DataQuery(
                    job_id=job_id,
                    tenant_id=tenant_id,
                    question=question,
                    image_id=None,
                    run_ids=run_ids,
                )
            )
            limitations.extend(result.limitations)
            for item in result.evidence:
                fingerprint = json.dumps(
                    item.model_dump(mode="json"),
                    ensure_ascii=False,
                    sort_keys=True,
                    separators=(",", ":"),
                )
                if fingerprint in seen:
                    continue
                seen.add(fingerprint)
                evidence.append(item)
        return tuple(evidence), tuple(dict.fromkeys(limitations))

    def _technical_summary(
        self,
        *,
        snapshot: JobExportSnapshot,
        runs: Sequence[SegmentationRunDTO],
        evidence: Sequence[ToolEvidence],
    ) -> tuple[str, str, str | None, bool]:
        fallback = _deterministic_summary(snapshot, runs, evidence)
        provider = self.llm_provider
        if provider is None or not evidence:
            return fallback, "deterministic_fallback", None, True
        try:
            if provider.health().status == "unavailable":
                return fallback, "deterministic_fallback", None, True
            generated = provider.generate_conversation(
                question=(
                    "只写一个中文句子，必须忠实包含 DATA_EVIDENCE 的质量解释、"
                    "视野限制和一个建议步骤；不得出现任何实验数值、百分比、单位、"
                    "参数或运行编号；句末必须保留 [D1]。"
                    "不要声称元素、晶相、价态、性能或因果机理。"
                ),
                query_type="analysis_data",
                history=(),
                data_evidence=evidence[:1],
                contexts=(),
                material_context=None,
            )
            validate_conversation_answer(
                answer=generated.answer,
                limitations=generated.limitations,
                used_data_ids=generated.used_data_ids,
                used_citation_ids=generated.used_citation_ids,
                data_evidence=evidence[:1],
                citation_contexts=(),
                allow_uncited_general_chat=False,
            )
            if not generated.used_data_ids:
                raise CitationValidationError(
                    "report synthesis must cite at least one data evidence item"
                )
            without_marker = re.sub(r"\[D\d+\]", "", generated.answer)
            if re.search(r"\d", without_marker):
                raise CitationValidationError(
                    "report qualitative synthesis must not contain numerical claims"
                )
            if any(term not in generated.answer for term in ("人工", "视野", "建议")):
                raise CitationValidationError(
                    "report qualitative synthesis omitted required review guardrails"
                )
            return (
                f"{fallback}\n\n本地模型综合：{generated.answer}",
                "local_llm",
                provider.model,
                False,
            )
        except (
            AnswerProviderError,
            CitationValidationError,
            TypeError,
            ValueError,
        ) as error:
            logger.info(
                "scientific_report_llm_fallback",
                extra={"reason": type(error).__name__},
            )
            return fallback, "deterministic_fallback", provider.model, True

    def _store_report(self, job_id: str, filename: str, data: bytes) -> StoredFile:
        path = self.file_store.paths.report_file(job_id, filename)
        self.file_store.atomic_write_bytes(path, data)
        return StoredFile(
            path=path,
            relative_path=self.file_store.paths.relative_path(path),
            filename=filename,
            size_bytes=len(data),
            sha256=hashlib.sha256(data).hexdigest(),
        )


def _build_preview(
    *,
    report_id: str,
    snapshot: JobExportSnapshot,
    runs: Sequence[SegmentationRunDTO],
    generated_at: datetime,
    technical_summary: str,
    synthesis_provider: str,
    synthesis_model: str | None,
    fallback_used: bool,
    evidence: Sequence[ToolEvidence],
    tool_limitations: Sequence[str],
) -> ScientificReportPreviewDTO:
    quality_status = _overall_quality(runs)
    image_count = len({run.image_id for run in runs})
    analysis_mode = "batch" if image_count > 1 else "single_image"
    scales = [run.configuration.scale_nm_per_pixel for run in runs]
    scale_status = (
        "physical"
        if all(value is not None for value in scales)
        else "pixel_only"
        if all(value is None for value in scales)
        else "mixed"
    )
    evidence_ids = [f"D{index}" for index in range(1, len(evidence) + 1)]
    images_by_id = {image.image_id: image for image in snapshot.images}
    run_summaries = [_run_summary(run, image=images_by_id.get(run.image_id)) for run in runs]
    batch_summary = _batch_summary(runs) if analysis_mode == "batch" else None
    # Report synthesis currently uses only run-scoped data-tool evidence.  Do not
    # silently promote unrelated citations from earlier task conversations into
    # the report; literature citations can be added once the report workflow
    # performs its own managed-knowledge retrieval.
    knowledge_citations: list[Citation] = []
    limitations = _limitations(
        runs,
        scale_status=scale_status,
        tool_limitations=tool_limitations,
        knowledge_citations=knowledge_citations,
    )
    return ScientificReportPreviewDTO(
        report_id=report_id,
        job_id=snapshot.job.job_id,
        title=(
            f"{snapshot.job.name} - SEM 纳米颗粒批量分析报告"
            if analysis_mode == "batch"
            else f"{snapshot.job.name} - SEM 纳米颗粒分析报告"
        ),
        generated_at=generated_at,
        selected_run_ids=[run.run_id for run in runs],
        analysis_mode=analysis_mode,
        quality_status=quality_status,
        scale_status=scale_status,
        technical_summary=technical_summary,
        synthesis_provider=synthesis_provider,
        synthesis_model=synthesis_model,
        fallback_used=fallback_used,
        headline_metrics=_headline_metrics(runs, batch_summary=batch_summary),
        findings=_findings(runs, scale_status=scale_status, evidence_ids=evidence_ids),
        run_summaries=run_summaries,
        batch_summary=batch_summary,
        recommendations=_recommendations(runs, scale_status=scale_status),
        methodology=[
            (
                f"分析对象为所选的 {len(runs)} 个终态运行，覆盖 {image_count} 个图像视野；"
                "每个运行冻结原图 SHA-256、ROI、模型版本、"
                "推理参数、后处理与质量门控配置。"
            ),
            (
                "颗粒数量、面积、等效粒径、覆盖率和密度由 canonical 实例与形貌统计代码计算；"
                f"报告层不在浏览器或大模型中重算数值。{_scale_methodology(runs)}"
            ),
            (
                "批量模式按运行汇总均值、标准差、四分位数、极值和变异系数，"
                "并用 1.5×IQR 规则标记待复核视野。"
                if analysis_mode == "batch"
                else "单图模式保留当前运行的完整形貌指标与质量状态。"
            ),
            (
                "确定性数据工具分别执行任务概览、粒径分布、异常与复核查询，"
                "并把来源运行和单位作为证据保留。"
            ),
            (
                "本地大模型只对已验证证据做一次语言综合；模型不可用或证据标记校验失败时，"
                "自动采用确定性摘要。"
            ),
        ],
        limitations=limitations,
        further_questions=_further_questions(
            runs,
            scale_status=scale_status,
            knowledge_citations=knowledge_citations,
        ),
        data_evidence=list(evidence),
        knowledge_citations=knowledge_citations,
        provenance=[_provenance(run) for run in runs],
    )


def _report_interpretation_evidence(
    runs: Sequence[SegmentationRunDTO],
) -> ToolEvidence:
    quality = _overall_quality(runs)
    image_count = len({run.image_id for run in runs})
    quality_interpretation = (
        "自动质量门控通过，但不替代人工边界抽查"
        if quality is QualityStatus.PASS
        else "自动质量门控存在告警，应先人工复核实例边界与告警区域"
    )
    return ToolEvidence(
        tool_name="interpret_report_guardrails",
        validated_arguments={
            "run_ids": [run.run_id for run in runs],
            "intent": "qualitative_report_synthesis",
        },
        aggregates={
            "quality_interpretation": quality_interpretation,
            "scope_limitation": (
                "所选多个图像视野仍不能自动代表完整样品"
                if image_count > 1
                else "单个图像视野不能代表完整样品"
            ),
            "allowed_next_steps": [
                "抽查小颗粒、粘连区和边界",
                "增加重复视野",
                "增加对照运行或人工真值",
            ],
        },
        source_run_ids=[run.run_id for run in runs],
        quality_warnings=[
            reason
            for run in runs
            for reason in (run.quality.reasons if run.quality is not None else [])
        ],
    )


def _report_synthesis_evidence(
    runs: Sequence[SegmentationRunDTO],
) -> ToolEvidence:
    if len({run.image_id for run in runs}) > 1:
        batch = _batch_summary(runs)
        distributions = {item.key: item for item in batch.distributions}
        display_metrics = {
            "image_count": f"{batch.image_count} 张",
            "run_count": f"{batch.run_count} 个",
            "total_particle_count": f"{batch.total_particle_count} 个",
            "mean_particle_count": _display_distribution_mean(distributions["particle_count"]),
            "mean_coverage": _display_distribution_mean(distributions["coverage"]),
            "quality_pass_count": batch.quality_pass_count,
            "quality_warning_count": batch.quality_warning_count,
            "quality_review_count": batch.quality_review_count,
            "outlier_count": len(batch.outliers),
        }
        diameter = distributions.get("mean_equivalent_diameter")
        if diameter is not None:
            display_metrics["mean_equivalent_diameter"] = _display_distribution_mean(diameter)
        return ToolEvidence(
            tool_name="build_batch_report_synthesis_snapshot",
            validated_arguments={
                "run_ids": [run.run_id for run in runs],
                "intent": "scientific_batch_report_synthesis",
            },
            aggregates={
                "display_metrics": display_metrics,
                "distribution_statistics": [
                    item.model_dump(mode="json") for item in batch.distributions
                ],
                "outliers": [item.model_dump(mode="json") for item in batch.outliers],
            },
            units={item.key: item.unit for item in batch.distributions},
            source_run_ids=[run.run_id for run in runs],
            quality_warnings=[
                reason
                for run in runs
                for reason in (run.quality.reasons if run.quality is not None else [])
            ],
        )

    primary = runs[0]
    summary = primary.summary
    quality = primary.quality
    if summary is None or quality is None:
        raise ValueError("report synthesis evidence requires summary and quality")
    scale = primary.configuration.scale_nm_per_pixel
    roi_area = (
        f"{_number(summary.roi_area_px * (scale**2) / 1_000_000)} µm²"
        if scale is not None
        else f"{_number(summary.roi_area_px, digits=0)} px²"
    )
    diameter_display = (
        f"{_number(summary.mean_equivalent_diameter_nm)} nm"
        if summary.mean_equivalent_diameter_nm is not None
        else f"{_number(summary.mean_equivalent_diameter_px)} px"
    )
    density_display = (
        f"{_number(summary.number_density_um2)} µm⁻²"
        if summary.number_density_um2 is not None
        else f"{_number(summary.number_density_px2, digits=6)} px⁻²"
    )
    perimeter_density_display = (
        f"{_number(summary.perimeter_density_um)} µm⁻¹"
        if summary.perimeter_density_um is not None
        else f"{_number(summary.perimeter_density_px, digits=6)} px⁻¹"
    )
    scale_display = f"{_number(scale, digits=4)} nm/px" if scale is not None else "仅像素尺度"
    return ToolEvidence(
        tool_name="build_report_synthesis_snapshot",
        validated_arguments={
            "run_ids": [run.run_id for run in runs],
            "intent": "scientific_report_synthesis",
        },
        aggregates={
            "display_metrics": {
                "particle_count": f"{_number(summary.particle_count, digits=0)} 个",
                "roi_area": roi_area,
                "mean_equivalent_diameter": diameter_display,
                "coverage": f"{_number(summary.coverage_ratio * 100)}%",
                "number_density": density_display,
                "perimeter_density": perimeter_density_display,
                "scale": scale_display,
                "quality_status": quality.status.value,
            },
            "selected_run_count": len(runs),
        },
        units={
            "particle_count": "count",
            "roi_area": "µm²" if scale is not None else "px²",
            "mean_equivalent_diameter": (
                "nm" if summary.mean_equivalent_diameter_nm is not None else "px"
            ),
            "coverage": "%",
            "number_density": "µm⁻²" if summary.number_density_um2 is not None else "px⁻²",
            "perimeter_density": ("µm⁻¹" if summary.perimeter_density_um is not None else "px⁻¹"),
            "scale": "nm/px" if scale is not None else "pixel_only",
        },
        source_run_ids=[run.run_id for run in runs],
        quality_warnings=[
            reason
            for run in runs
            for reason in (run.quality.reasons if run.quality is not None else [])
        ],
    )


def _display_distribution_mean(distribution: BatchMetricDistributionDTO) -> str:
    suffix = "%" if distribution.unit == "%" else f" {distribution.unit}"
    return f"{_number(distribution.mean)}{suffix}"


def _scale_methodology(runs: Sequence[SegmentationRunDTO]) -> str:
    calibrated = [
        run.configuration.scale_calibration
        for run in runs
        if run.configuration.scale_calibration is not None
    ]
    if calibrated:
        bases = list(
            dict.fromkeys(
                (
                    f"{_number(item.physical_length_nm)} nm / "
                    f"{_number(item.pixel_length_px)} px"
                    + (f"（原图标签：{item.label_text}）" if item.label_text else "")
                )
                for item in calibrated
            )
        )
        return f" 物理尺度来自已冻结的可见标尺复核：{'；'.join(bases)}。"
    if all(run.configuration.scale_nm_per_pixel is not None for run in runs):
        return " 物理尺度已冻结在运行配置中，但当前快照未记录可见标尺测量明细。"
    return " 缺少已冻结物理尺度的运行保留像素单位，不进行推测换算。"


def _headline_metrics(
    runs: Sequence[SegmentationRunDTO],
    *,
    batch_summary: BatchResultSummaryDTO | None,
) -> list[ReportMetricDTO]:
    if batch_summary is not None:
        source_run_ids = [run.run_id for run in runs]
        distributions = {item.key: item for item in batch_summary.distributions}
        count_distribution = distributions["particle_count"]
        coverage_distribution = distributions["coverage"]
        diameter_distribution = distributions.get("mean_equivalent_diameter")
        metrics = [
            ReportMetricDTO(
                key="image_count",
                label="图像视野",
                display_value=_number(batch_summary.image_count, digits=0),
                unit="张",
                definition="批量报告覆盖的不同 SEM 图像视野数量。",
                source_run_ids=source_run_ids,
            ),
            ReportMetricDTO(
                key="run_count",
                label="终态运行",
                display_value=_number(batch_summary.run_count, digits=0),
                unit="个",
                definition="进入本次批量汇总且具有完整质量与形貌统计的运行数量。",
                source_run_ids=source_run_ids,
            ),
            ReportMetricDTO(
                key="total_particle_count",
                label="颗粒总数",
                display_value=_number(batch_summary.total_particle_count, digits=0),
                unit="个",
                definition="所有所选运行 canonical 实例数量之和；多模型运行会分别计入。",
                source_run_ids=source_run_ids,
            ),
            ReportMetricDTO(
                key="mean_particle_count",
                label="平均每运行颗粒数",
                display_value=_number(count_distribution.mean),
                unit="个",
                definition="运行级颗粒数量的算术平均，批间离散度见分布统计。",
                source_run_ids=source_run_ids,
            ),
            ReportMetricDTO(
                key="mean_coverage",
                label="平均覆盖率",
                display_value=f"{_number(coverage_distribution.mean)}%",
                unit="%",
                definition="运行级面积覆盖率的算术平均。",
                source_run_ids=source_run_ids,
            ),
        ]
        if diameter_distribution is not None:
            metrics.append(
                ReportMetricDTO(
                    key="mean_diameter",
                    label="批量平均等效粒径",
                    display_value=(
                        f"{_number(diameter_distribution.mean)} {diameter_distribution.unit}"
                    ),
                    unit=diameter_distribution.unit,
                    definition="各运行平均等效粒径的批量均值；四分位数和变异系数用于审查异质性。",
                    source_run_ids=source_run_ids,
                )
            )
        else:
            metrics.append(
                ReportMetricDTO(
                    key="quality_pass_count",
                    label="质量通过运行",
                    display_value=_number(batch_summary.quality_pass_count, digits=0),
                    unit="个",
                    definition="未触发自动质量警告或人工复核条件的终态运行数量。",
                    source_run_ids=source_run_ids,
                )
            )
        return metrics

    primary = runs[0]
    summary = primary.summary
    scale = primary.configuration.scale_nm_per_pixel
    if summary is None:
        return []
    roi_area_um2 = summary.roi_area_px * (scale**2) / 1_000_000 if scale is not None else None
    return [
        ReportMetricDTO(
            key="particle_count",
            label="颗粒数量",
            display_value=_number(summary.particle_count, digits=0),
            unit="个",
            definition="所选 ROI 内 canonical 后处理实例的数量。",
            source_run_ids=[primary.run_id],
        ),
        ReportMetricDTO(
            key="roi_area",
            label="有效 ROI 面积",
            display_value=(
                f"{_number(roi_area_um2)} µm²"
                if roi_area_um2 is not None
                else f"{_number(summary.roi_area_px, digits=0)} px²"
            ),
            unit="µm²" if roi_area_um2 is not None else "px²",
            definition="排除无效区域后参与统计的有效面积。",
            source_run_ids=[primary.run_id],
        ),
        ReportMetricDTO(
            key="mean_diameter",
            label="平均等效粒径",
            display_value=(
                f"{_number(summary.mean_equivalent_diameter_nm)} nm"
                if summary.mean_equivalent_diameter_nm is not None
                else f"{_number(summary.mean_equivalent_diameter_px)} px"
            ),
            unit="nm" if summary.mean_equivalent_diameter_nm is not None else "px",
            definition="与实例面积相同的圆的直径，按实例取算术平均。",
            source_run_ids=[primary.run_id],
        ),
        ReportMetricDTO(
            key="coverage",
            label="面积覆盖率",
            display_value=f"{_number(summary.coverage_ratio * 100)}%",
            unit="%",
            definition="前景实例面积占有效 ROI 面积的比例。",
            source_run_ids=[primary.run_id],
        ),
        ReportMetricDTO(
            key="number_density",
            label="颗粒数密度",
            display_value=(
                f"{_number(summary.number_density_um2)} µm⁻²"
                if summary.number_density_um2 is not None
                else f"{_number(summary.number_density_px2, digits=6)} px⁻²"
            ),
            unit="µm⁻²" if summary.number_density_um2 is not None else "px⁻²",
            definition="颗粒数量除以有效 ROI 面积。",
            source_run_ids=[primary.run_id],
        ),
        ReportMetricDTO(
            key="perimeter_density",
            label="周长密度",
            display_value=(
                f"{_number(summary.perimeter_density_um)} µm⁻¹"
                if summary.perimeter_density_um is not None
                else f"{_number(summary.perimeter_density_px, digits=6)} px⁻¹"
            ),
            unit="µm⁻¹" if summary.perimeter_density_um is not None else "px⁻¹",
            definition="全部实例周长之和除以有效 ROI 面积。",
            source_run_ids=[primary.run_id],
        ),
    ]


def _run_summary(run: SegmentationRunDTO, *, image: Any | None) -> ReportRunSummaryDTO:
    summary = run.summary
    quality = run.quality
    scale = run.configuration.scale_nm_per_pixel
    if summary is None or quality is None:
        raise ValueError("report run is missing summary or quality")
    roi_area_um2 = summary.roi_area_px * (scale**2) / 1_000_000 if scale else None
    return ReportRunSummaryDTO(
        run_id=run.run_id,
        image_id=run.image_id,
        filename=getattr(image, "filename", None),
        sample_id=getattr(image, "sample_id", None),
        model_id=run.model_id,
        quality_status=quality.status,
        scale_nm_per_pixel=scale,
        particle_count=summary.particle_count,
        roi_area=(
            f"{_number(roi_area_um2)} µm²"
            if roi_area_um2 is not None
            else f"{_number(summary.roi_area_px, digits=0)} px²"
        ),
        number_density=(
            f"{_number(summary.number_density_um2)} µm⁻²"
            if summary.number_density_um2 is not None
            else f"{_number(summary.number_density_px2, digits=6)} px⁻²"
        ),
        mean_equivalent_diameter=(
            f"{_number(summary.mean_equivalent_diameter_nm)} nm"
            if summary.mean_equivalent_diameter_nm is not None
            else f"{_number(summary.mean_equivalent_diameter_px)} px"
        ),
        coverage=f"{_number(summary.coverage_ratio * 100)}%",
        perimeter_density=(
            f"{_number(summary.perimeter_density_um)} µm⁻¹"
            if summary.perimeter_density_um is not None
            else f"{_number(summary.perimeter_density_px, digits=6)} px⁻¹"
        ),
        quality_reasons=list(quality.reasons),
    )


def _batch_summary(runs: Sequence[SegmentationRunDTO]) -> BatchResultSummaryDTO:
    metric_rows = _batch_metric_rows(runs)
    distributions: list[BatchMetricDistributionDTO] = []
    outliers: list[BatchOutlierDTO] = []
    for key, label, unit, rows in metric_rows:
        values = [value for _, value in rows]
        if not values:
            continue
        q1 = _percentile(values, 0.25)
        median = _percentile(values, 0.5)
        q3 = _percentile(values, 0.75)
        mean = fmean(values)
        std_dev = pstdev(values)
        distribution = BatchMetricDistributionDTO(
            key=key,
            label=label,
            unit=unit,
            sample_count=len(values),
            mean=mean,
            std_dev=std_dev,
            minimum=min(values),
            q1=q1,
            median=median,
            q3=q3,
            maximum=max(values),
            coefficient_of_variation=(std_dev / abs(mean) if mean else None),
        )
        distributions.append(distribution)
        if len(values) < 4:
            continue
        spread = q3 - q1
        low = q1 - 1.5 * spread
        high = q3 + 1.5 * spread
        for run, value in rows:
            direction = "low" if value < low else "high" if value > high else None
            if direction is None:
                continue
            outliers.append(
                BatchOutlierDTO(
                    run_id=run.run_id,
                    image_id=run.image_id,
                    metric_key=key,
                    metric_label=label,
                    value=value,
                    unit=unit,
                    direction=direction,
                )
            )

    quality_statuses = [run.quality.status for run in runs if run.quality is not None]
    return BatchResultSummaryDTO(
        image_count=len({run.image_id for run in runs}),
        run_count=len(runs),
        model_count=len({run.model_id for run in runs}),
        total_particle_count=sum(
            run.summary.particle_count for run in runs if run.summary is not None
        ),
        quality_pass_count=quality_statuses.count(QualityStatus.PASS),
        quality_warning_count=quality_statuses.count(QualityStatus.WARN),
        quality_review_count=quality_statuses.count(QualityStatus.REVIEW_REQUIRED),
        distributions=distributions,
        outliers=outliers[:20],
    )


def _batch_metric_rows(
    runs: Sequence[SegmentationRunDTO],
) -> list[tuple[str, str, str, list[tuple[SegmentationRunDTO, float]]]]:
    summaries = [(run, run.summary) for run in runs if run.summary is not None]
    physical_diameter = all(
        summary.mean_equivalent_diameter_nm is not None for _, summary in summaries
    )
    physical_density = all(summary.number_density_um2 is not None for _, summary in summaries)
    physical_perimeter = all(summary.perimeter_density_um is not None for _, summary in summaries)
    return [
        (
            "particle_count",
            "颗粒数量",
            "个",
            [(run, float(summary.particle_count)) for run, summary in summaries],
        ),
        (
            "coverage",
            "面积覆盖率",
            "%",
            [(run, summary.coverage_ratio * 100.0) for run, summary in summaries],
        ),
        (
            "mean_equivalent_diameter",
            "平均等效粒径",
            "nm" if physical_diameter else "px",
            [
                (
                    run,
                    _required_float(
                        summary.mean_equivalent_diameter_nm
                        if physical_diameter
                        else summary.mean_equivalent_diameter_px
                    ),
                )
                for run, summary in summaries
                if (
                    summary.mean_equivalent_diameter_nm
                    if physical_diameter
                    else summary.mean_equivalent_diameter_px
                )
                is not None
            ],
        ),
        (
            "number_density",
            "颗粒数密度",
            "µm⁻²" if physical_density else "px⁻²",
            [
                (
                    run,
                    _required_float(
                        summary.number_density_um2
                        if physical_density
                        else summary.number_density_px2
                    ),
                )
                for run, summary in summaries
            ],
        ),
        (
            "perimeter_density",
            "周长密度",
            "µm⁻¹" if physical_perimeter else "px⁻¹",
            [
                (
                    run,
                    _required_float(
                        summary.perimeter_density_um
                        if physical_perimeter
                        else summary.perimeter_density_px
                    ),
                )
                for run, summary in summaries
            ],
        ),
    ]


def _percentile(values: Sequence[float], quantile: float) -> float:
    ordered = sorted(values)
    if not ordered:
        raise ValueError("percentile requires at least one value")
    position = (len(ordered) - 1) * quantile
    lower = int(position)
    upper = min(lower + 1, len(ordered) - 1)
    weight = position - lower
    return ordered[lower] * (1.0 - weight) + ordered[upper] * weight


def _required_float(value: float | None) -> float:
    if value is None:
        raise ValueError("report metric unexpectedly missing")
    return float(value)


def _findings(
    runs: Sequence[SegmentationRunDTO],
    *,
    scale_status: str,
    evidence_ids: Sequence[str],
) -> list[ReportFindingDTO]:
    findings: list[ReportFindingDTO] = []
    overall = _overall_quality(runs)
    reasons = list(
        dict.fromkeys(
            reason for run in runs for reason in (run.quality.reasons if run.quality else [])
        )
    )
    if overall is QualityStatus.PASS:
        findings.append(
            ReportFindingDTO(
                title="所选运行通过当前质量门控",
                interpretation=(
                    "当前配置未触发已定义的自动风险阈值；这支持继续解释形貌统计，"
                    "但不替代人工边界抽查或外部真值验证。"
                ),
                severity="info",
                evidence_ids=list(evidence_ids[:1]),
                source_run_ids=[run.run_id for run in runs],
            )
        )
    else:
        findings.append(
            ReportFindingDTO(
                title="质量门控要求先完成针对性复核",
                interpretation=(
                    "触发项：" + "、".join(_human_reason(reason) for reason in reasons)
                    if reasons
                    else "至少一个所选运行被标记为需要复核。"
                ),
                severity="review" if overall is QualityStatus.REVIEW_REQUIRED else "caution",
                evidence_ids=list(evidence_ids[:1]),
                source_run_ids=[run.run_id for run in runs],
            )
        )

    if scale_status == "physical":
        findings.append(
            ReportFindingDTO(
                title="全部所选运行具备物理尺度",
                interpretation=(
                    "粒径、面积和密度已优先转换为 nm/µm 单位，可在相同尺度定义下解释；"
                    "报告仍保留像素级来源用于复现。"
                ),
                severity="info",
                evidence_ids=list(evidence_ids[:1]),
                source_run_ids=[run.run_id for run in runs],
            )
        )
    else:
        findings.append(
            ReportFindingDTO(
                title="尺度信息不足限制了跨图和材料层解释",
                interpretation=(
                    "至少一个运行只能给出 px、px² 或 px⁻¹；这些值适合算法复核，"
                    "不能直接当作真实粒径或物理密度。"
                ),
                severity="review" if scale_status == "pixel_only" else "caution",
                evidence_ids=list(evidence_ids[:1]),
                source_run_ids=[run.run_id for run in runs],
            )
        )

    primary = runs[0]
    summary = primary.summary
    image_count = len({run.image_id for run in runs})
    if summary is not None and image_count == 1:
        findings.append(
            ReportFindingDTO(
                title=(
                    f"主运行识别 {_number(summary.particle_count, digits=0)} 个实例，"
                    f"覆盖率为 {_number(summary.coverage_ratio * 100)}%"
                ),
                interpretation=(
                    "该结果描述所选 ROI 与当前模型/阈值下的分割形貌，"
                    "应结合叠加图检查粘连、边界和小颗粒是否被一致处理。"
                ),
                severity="info",
                evidence_ids=list(evidence_ids[:2]),
                source_run_ids=[primary.run_id],
            )
        )

    if image_count > 1:
        batch = _batch_summary(runs)
        distributions = {item.key: item for item in batch.distributions}
        count_distribution = distributions["particle_count"]
        coverage_distribution = distributions["coverage"]
        largest_cv = max(
            (item for item in batch.distributions if item.coefficient_of_variation is not None),
            key=lambda item: item.coefficient_of_variation or 0.0,
            default=None,
        )
        findings.append(
            ReportFindingDTO(
                title=(
                    f"{batch.image_count} 个视野共识别 "
                    f"{_number(batch.total_particle_count, digits=0)} 个实例"
                ),
                interpretation=(
                    f"单运行颗粒数中位数为 {_number(count_distribution.median)}，"
                    f"覆盖率四分位区间为 {_number(coverage_distribution.q1)}%–"
                    f"{_number(coverage_distribution.q3)}%。"
                    + (
                        f"离散度最高的指标为{largest_cv.label}"
                        f"（CV={_number((largest_cv.coefficient_of_variation or 0) * 100)}%）。"
                        if largest_cv is not None
                        else ""
                    )
                ),
                severity="caution" if batch.outliers else "info",
                evidence_ids=list(evidence_ids),
                source_run_ids=[run.run_id for run in runs],
            )
        )
        findings.append(
            ReportFindingDTO(
                title=(
                    f"批量异常筛查标记 {len(batch.outliers)} 个指标-运行组合"
                    if batch.outliers
                    else "批量异常筛查未发现 1.5×IQR 极端视野"
                ),
                interpretation=(
                    "异常标记用于定位需要回看原图、掩码和尺度信息的视野，"
                    "不等同于自动删除数据或判断样品异常。"
                ),
                severity="review" if batch.outliers else "info",
                evidence_ids=list(evidence_ids),
                source_run_ids=list(dict.fromkeys(item.run_id for item in batch.outliers)),
            )
        )
    elif len(runs) > 1:
        counts = [run.summary.particle_count for run in runs if run.summary is not None]
        coverages = [run.summary.coverage_ratio for run in runs if run.summary is not None]
        count_span = max(counts) - min(counts)
        coverage_span = (max(coverages) - min(coverages)) * 100
        findings.append(
            ReportFindingDTO(
                title="不同运行对同一形貌的分割口径存在差异",
                interpretation=(
                    f"所选运行的颗粒数范围相差 {_number(count_span, digits=0)}，"
                    f"覆盖率极差为 {_number(coverage_span)} 个百分点。"
                    "差异只能说明模型/参数敏感性，不能据此自动判定最佳模型。"
                ),
                severity="caution" if count_span or coverage_span else "info",
                evidence_ids=list(evidence_ids),
                source_run_ids=[run.run_id for run in runs],
            )
        )
    return findings


def _recommendations(
    runs: Sequence[SegmentationRunDTO],
    *,
    scale_status: str,
) -> list[ReportRecommendationDTO]:
    recommendations: list[ReportRecommendationDTO] = []
    all_run_ids = [run.run_id for run in runs]
    if scale_status != "physical":
        recommendations.append(
            ReportRecommendationDTO(
                priority=1,
                action="补录并核验 scale_nm_per_pixel，再生成物理单位报告。",
                rationale=(
                    "缺少比例尺时，粒径和密度只能停留在像素坐标，"
                    "无法支持跨倍率、跨图像或跨样品比较。"
                ),
                verification=(
                    "用图中比例尺或仪器元数据独立换算一次；重新运行后确认报告同时出现 nm、µm²、"
                    "µm⁻² 与 µm⁻¹。"
                ),
                source_run_ids=all_run_ids,
            )
        )

    for run in runs:
        quality = run.quality
        if quality is None:
            continue
        reason_text = " ".join(quality.reasons).casefold()
        if "small_fragment" in reason_text or "fragment" in reason_text:
            recommendations.append(
                ReportRecommendationDTO(
                    priority=2,
                    action=(
                        "在实例编号图中定位碎片化小连通域，并从当前 "
                        f"min_area_px={run.inference.min_area_px} 开始做一次单变量上调复核。"
                    ),
                    rationale="小碎片比例偏高会抬高颗粒数并压低平均等效粒径。",
                    verification=(
                        "比较父子运行的颗粒数、粒径分布和覆盖率，同时确认真实小颗粒没有被整体删除。"
                    ),
                    source_run_ids=[run.run_id],
                )
            )
        if "foreground" in reason_text or "coverage" in reason_text:
            threshold = run.inference.threshold
            recommendations.append(
                ReportRecommendationDTO(
                    priority=2,
                    action=(
                        "保持 ROI、模型和其他后处理不变，仅围绕当前阈值"
                        f"{f' {threshold:g}' if threshold is not None else ''}建立复核子运行。"
                    ),
                    rationale="前景占比异常常与阈值、成像对比度或背景误分割有关。",
                    verification="叠加图中的颗粒边界更贴合原图，且覆盖率回到可解释范围。",
                    source_run_ids=[run.run_id],
                )
            )
        if "border" in reason_text or "edge" in reason_text:
            recommendations.append(
                ReportRecommendationDTO(
                    priority=2,
                    action="复核 ROI 边界与 exclude_border 设置，单独检查截断颗粒。",
                    rationale="边界截断会同时影响数量、面积、周长和粒径统计。",
                    verification="父子运行采用明确一致的边界规则，报告中不再混合保留与排除口径。",
                    source_run_ids=[run.run_id],
                )
            )
        if "agglomer" in reason_text or "merged" in reason_text:
            recommendations.append(
                ReportRecommendationDTO(
                    priority=2,
                    action="在粘连区域核对实例编号，并单独测试 watershed 开关。",
                    rationale="团聚颗粒被合并会降低计数、抬高等效粒径并改变周长密度。",
                    verification="对同一批粘连区域人工点数，确认子运行的拆分方向更接近人工判读。",
                    source_run_ids=[run.run_id],
                )
            )

    recommendations.append(
        ReportRecommendationDTO(
            priority=3,
            action="按实例编号抽查小颗粒、粘连区和 ROI 边界，并记录漏检/误检类型。",
            rationale="自动质量门控只能覆盖预定义异常，不能证明每个实例边界正确。",
            verification="形成可追溯的人工复核记录；若有修正，使用修正掩码创建不可变子运行。",
            source_run_ids=all_run_ids,
        )
    )
    image_count = len({run.image_id for run in runs})
    if image_count > 1:
        batch = _batch_summary(runs)
        recommendations.append(
            ReportRecommendationDTO(
                priority=4,
                action="优先回看批量 IQR 筛查标记的视野，并核对倍率、尺度、成像质量与分割边界。",
                rationale=(
                    "跨视野极端值可能来自真实空间异质性，也可能来自成像或分割口径差异，"
                    "不能在未复核时自动剔除。"
                ),
                verification=(
                    f"逐一审查 {len(batch.outliers)} 个异常指标-运行组合，"
                    "记录保留、修正或排除理由。"
                ),
                source_run_ids=list(dict.fromkeys(item.run_id for item in batch.outliers))
                or all_run_ids,
            )
        )
        recommendations.append(
            ReportRecommendationDTO(
                priority=5,
                action="按样品与实验条件增加独立重复视野，并预先固定抽样规则。",
                rationale="当前批量汇总量化了所选视野的变异，但不能自动证明抽样具有代表性。",
                verification="新增视野后重新生成报告，比较均值、四分位区间和 CV 是否趋于稳定。",
                source_run_ids=all_run_ids,
            )
        )
    elif len(runs) > 1:
        recommendations.append(
            ReportRecommendationDTO(
                priority=4,
                action="使用共同授权的像素或实例级 GT 决定模型，而不是选择颗粒数更高的运行。",
                rationale="模型间统计差异反映口径敏感性，不等价于科学准确率。",
                verification="在固定独立集上报告明确的像素级和实例级指标及容差。",
                source_run_ids=all_run_ids,
            )
        )
    else:
        recommendations.append(
            ReportRecommendationDTO(
                priority=4,
                action="若结论将用于样品比较，为同一图像保留至少一个对照运行或人工 GT。",
                rationale="单运行没有模型/参数敏感性基线，难以判断结论是否依赖当前配置。",
                verification="报告能够显示同 ROI 的可比运行，并明确差异是否超过人工容差。",
                source_run_ids=all_run_ids,
            )
        )
    deduplicated: list[ReportRecommendationDTO] = []
    seen: set[str] = set()
    for item in recommendations:
        if item.action in seen:
            continue
        seen.add(item.action)
        deduplicated.append(item)
    return [
        item.model_copy(update={"priority": index})
        for index, item in enumerate(deduplicated, start=1)
    ]


def _limitations(
    runs: Sequence[SegmentationRunDTO],
    *,
    scale_status: str,
    tool_limitations: Sequence[str],
    knowledge_citations: Sequence[Citation],
) -> list[str]:
    limitations = [
        "分割与形貌统计属于描述性结果，不单独证明元素组成、价态、晶相、性能或因果机理。",
        "自动质量门控只检查已配置规则；未触发告警不等于像素级或实例级科学准确率已验证。",
    ]
    if scale_status != "physical":
        limitations.append("部分或全部运行缺少物理尺度，像素单位不能直接用于跨倍率比较。")
    if len({run.image_id for run in runs}) == 1:
        limitations.append("当前报告只覆盖一个图像视野，不能代表完整样品的空间异质性。")
    else:
        limitations.append(
            "批量统计只描述所选视野；若视野抽样、倍率或实验条件不一致，"
            "均值与离散度不能自动外推到完整样品。"
        )
        if len({run.model_id for run in runs}) > 1:
            limitations.append(
                "所选批量运行包含多个模型，跨视野差异与模型口径差异可能混杂；"
                "样品比较应优先使用同一冻结模型与参数。"
            )
    if any(run.quality and run.quality.status is not QualityStatus.PASS for run in runs):
        limitations.append("至少一个所选运行带质量告警，复核完成前不应把统计写成确定性材料结论。")
    if not knowledge_citations:
        limitations.append("本报告未使用受管文献证据，因此不解释材料机理或外部性能规律。")
    limitations.extend(
        limitation
        for limitation in tool_limitations
        if limitation and limitation not in limitations
    )
    return list(dict.fromkeys(limitations))


def _further_questions(
    runs: Sequence[SegmentationRunDTO],
    *,
    scale_status: str,
    knowledge_citations: Sequence[Citation],
) -> list[str]:
    questions: list[str] = []
    if scale_status != "physical":
        questions.append("图像比例尺能否从原始仪器元数据或可读比例尺中可靠恢复？")
    image_count = len({run.image_id for run in runs})
    if len(runs) == 1:
        questions.append("在同一 ROI 上更换模型或阈值后，主要统计是否保持在可接受容差内？")
    elif image_count > 1:
        questions.append("异常视野来自真实空间异质性、成像质量变化，还是分割口径变化？")
    if any(run.quality and run.quality.status is not QualityStatus.PASS for run in runs):
        questions.append("质量告警集中在哪些实例或区域，修正后会怎样改变统计分布？")
    if not knowledge_citations:
        questions.append("若要解释材料机理，哪些受管文献或实验表征应先导入知识库？")
    questions.append("当前视野与其他重复视野之间的变异是否大于模型/参数带来的变异？")
    return questions


def _provenance(run: SegmentationRunDTO) -> ReportProvenanceDTO:
    configuration = run.configuration
    return ReportProvenanceDTO(
        run_id=run.run_id,
        model_id=run.model_id,
        model_version=configuration.model_version,
        image_sha256=configuration.image_sha256,
        model_bundle_sha256=(
            configuration.model_bundle.bundle_id
            if configuration.model_bundle is not None
            else configuration.weight_sha256
        ),
        scale_nm_per_pixel=configuration.scale_nm_per_pixel,
        roi_mode=configuration.roi_mode.value,
        box_revision=configuration.box_revision,
        threshold=run.inference.threshold,
        min_area_px=run.inference.min_area_px,
    )


def _deterministic_summary(
    snapshot: JobExportSnapshot,
    runs: Sequence[SegmentationRunDTO],
    evidence: Sequence[ToolEvidence],
) -> str:
    if len({run.image_id for run in runs}) > 1:
        batch = _batch_summary(runs)
        distributions = {item.key: item for item in batch.distributions}
        counts = distributions["particle_count"]
        coverage = distributions["coverage"]
        diameter = distributions.get("mean_equivalent_diameter")
        marker = (
            " [D2]"
            if len(evidence) > 1 and evidence[0].tool_name == "interpret_report_guardrails"
            else " [D1]"
            if evidence
            else ""
        )
        diameter_sentence = (
            f"运行级平均等效粒径中位数为 {_number(diameter.median)} {diameter.unit}，"
            if diameter is not None
            else "当前批量运行未产生可用的等效粒径统计，"
        )
        quality_sentence = (
            f"{batch.quality_pass_count} 个运行通过质量门控，"
            f"{batch.quality_warning_count} 个带警告，"
            f"{batch.quality_review_count} 个需要复核。"
        )
        return (
            f"{snapshot.job.name} 的批量快照覆盖 {batch.image_count} 个图像视野、"
            f"{batch.run_count} 个终态运行，共识别 "
            f"{_number(batch.total_particle_count, digits=0)} 个颗粒；"
            f"单运行颗粒数中位数为 {_number(counts.median)}，"
            f"覆盖率中位数为 {_number(coverage.median)}%。{marker}\n\n"
            f"{diameter_sentence}{quality_sentence}"
            f"1.5×IQR 筛查标记 {len(batch.outliers)} 个待回看指标-运行组合。"
        )

    primary = runs[0]
    summary = primary.summary
    quality = primary.quality
    if summary is None or quality is None:
        return "所选运行尚无完整统计或质量报告，无法形成技术摘要。"
    marker = (
        " [D2]"
        if len(evidence) > 1 and evidence[0].tool_name == "interpret_report_guardrails"
        else " [D1]"
        if evidence
        else ""
    )
    scale = primary.configuration.scale_nm_per_pixel
    diameter_display = (
        f"{_number(summary.mean_equivalent_diameter_nm)} nm"
        if summary.mean_equivalent_diameter_nm is not None
        else f"{_number(summary.mean_equivalent_diameter_px)} px"
    )
    density_display = (
        f"{_number(summary.number_density_um2)} µm⁻²"
        if summary.number_density_um2 is not None
        else f"{_number(summary.number_density_px2, digits=6)} px⁻²"
    )
    scale_sentence = (
        f"图像尺度为 {_number(scale, digits=4)} nm/px，因此报告优先使用物理单位。"
        if scale is not None
        else "图像未提供可靠比例尺，因此粒径与密度只能使用像素单位，不能直接跨图比较。"
    )
    quality_sentence = (
        "当前运行通过自动质量门控，仍需人工抽查实例边界。"
        if quality.status is QualityStatus.PASS
        else "当前运行带质量告警，应先按报告中的可验证步骤创建复核子运行。"
    )
    return (
        f"{snapshot.job.name} 的主运行识别 {_number(summary.particle_count, digits=0)} 个颗粒，"
        f"平均等效粒径为 {diameter_display}，覆盖率为 "
        f"{_number(summary.coverage_ratio * 100)}%，颗粒数密度为 {density_display}。{marker}\n\n"
        f"{scale_sentence} {quality_sentence}"
    )


def _overall_quality(runs: Sequence[SegmentationRunDTO]) -> QualityStatus:
    statuses = [run.quality.status for run in runs if run.quality is not None]
    if QualityStatus.REVIEW_REQUIRED in statuses:
        return QualityStatus.REVIEW_REQUIRED
    if QualityStatus.WARN in statuses:
        return QualityStatus.WARN
    return QualityStatus.PASS


def _report_id(
    *,
    snapshot: JobExportSnapshot,
    run_ids: Sequence[str],
    generated_at: datetime,
    technical_summary: str,
    evidence: Sequence[ToolEvidence],
) -> str:
    payload = {
        "job_id": snapshot.job.job_id,
        "run_ids": list(run_ids),
        "generated_at": generated_at.isoformat(),
        "technical_summary": technical_summary,
        "evidence": [item.model_dump(mode="json") for item in evidence],
    }
    digest = hashlib.sha256(
        json.dumps(
            payload,
            ensure_ascii=False,
            sort_keys=True,
            separators=(",", ":"),
        ).encode("utf-8")
    ).hexdigest()
    return f"report_{digest[:24]}"


def _primary_overlay(
    snapshot: JobExportSnapshot,
    run_id: str,
    file_store: LocalFileStore,
) -> Path | None:
    for artifact_run_id, relative_path in snapshot.run_artifact_paths:
        if artifact_run_id != run_id or not relative_path.endswith("/overlay.png"):
            continue
        try:
            path = file_store.paths.require_managed(relative_path, must_exist=True)
        except (OSError, ValueError):
            return None
        if path.is_file():
            return path
    return None


def _render_docx(report: ScientificReportPreviewDTO, *, overlay: Path | None) -> bytes:
    document = Document()
    section = document.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.8)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.36)
    section.footer_distance = Inches(0.36)
    _configure_docx_styles(document)
    _set_docx_header_footer(section, report)

    kicker = document.add_paragraph()
    kicker.paragraph_format.space_after = Pt(3)
    _add_run(kicker, "NANOLOOP SCIENTIFIC REPORT", 9, color=_BLUE, bold=True)
    title = document.add_paragraph()
    title.paragraph_format.space_after = Pt(4)
    _add_run(title, report.title, 23, color=_INK, bold=True)
    subtitle = document.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(14)
    _add_run(
        subtitle,
        "可追溯 SEM 分割、形貌统计与质量复核技术报告",
        12,
        color=_MUTED,
    )
    for label, value in (
        ("报告 ID", report.report_id),
        ("生成时间", report.generated_at.strftime("%Y-%m-%d %H:%M:%S UTC")),
        ("运行范围", "、".join(report.selected_run_ids)),
        (
            "综合方式",
            (
                f"本地模型 {report.synthesis_model}"
                if report.synthesis_provider == "local_llm"
                else "确定性可信模板"
            ),
        ),
    ):
        paragraph = document.add_paragraph()
        paragraph.paragraph_format.space_after = Pt(2)
        _add_run(paragraph, f"{label}: ", 10.5, color=_INK, bold=True)
        _add_run(paragraph, value, 10.5, color=_INK)

    _add_heading(document, "技术摘要", level=1)
    if overlay is not None:
        _add_summary_with_overlay(document, report.technical_summary, overlay)
    else:
        _add_callout(document, report.technical_summary)

    _add_heading(document, "主要结果与单位定义", level=1)
    for finding in report.findings:
        paragraph = document.add_paragraph()
        paragraph.paragraph_format.space_after = Pt(4)
        _add_run(paragraph, finding.title, 11, color=_DARK_BLUE, bold=True)
        detail = document.add_paragraph(finding.interpretation)
        detail.paragraph_format.space_after = Pt(8)
        detail.paragraph_format.line_spacing = 1.1

    _add_metric_table(document, report)
    if report.batch_summary is not None:
        _add_heading(document, "批量分布与异常视野", level=2)
        _add_batch_distribution_table(document, report.batch_summary)

    _add_heading(document, "范围、数据与方法", level=1)
    for item in report.methodology:
        _add_list_item(document, item, numbered=False)

    _add_heading(document, "限制、不确定性与稳健性", level=1)
    for item in report.limitations:
        _add_list_item(document, item, numbered=False)

    _add_heading(document, "建议的下一步", level=1)
    for recommendation in report.recommendations:
        paragraph = _add_list_item(document, recommendation.action, numbered=True)
        paragraph.paragraph_format.keep_with_next = True
        rationale = document.add_paragraph()
        rationale.paragraph_format.left_indent = Inches(0.5)
        rationale.paragraph_format.space_after = Pt(2)
        _add_run(rationale, "依据: ", 10, color=_MUTED, bold=True)
        _add_run(rationale, recommendation.rationale, 10, color=_INK)
        verification = document.add_paragraph()
        verification.paragraph_format.left_indent = Inches(0.5)
        verification.paragraph_format.space_after = Pt(7)
        _add_run(verification, "验收: ", 10, color=_MUTED, bold=True)
        _add_run(verification, recommendation.verification, 10, color=_INK)

    _add_heading(document, "仍需回答的问题", level=1)
    for item in report.further_questions:
        _add_list_item(document, item, numbered=False)

    _add_heading(document, "证据与可追溯性", level=1)
    _add_provenance_table(document, report)
    if report.data_evidence:
        _add_heading(document, "确定性工具证据", level=2)
        for index, evidence_item in enumerate(report.data_evidence, start=1):
            paragraph = document.add_paragraph()
            paragraph.paragraph_format.space_after = Pt(5)
            _add_run(
                paragraph,
                f"[D{index}] {evidence_item.tool_name}",
                10,
                color=_DARK_BLUE,
                bold=True,
            )
            source_ids = "、".join(evidence_item.source_run_ids) or "无来源运行"
            _add_run(paragraph, f" - {source_ids}", 9.5, color=_MUTED)
    if report.knowledge_citations:
        _add_heading(document, "受管知识引用", level=2)
        for citation in report.knowledge_citations:
            _add_list_item(
                document,
                citation.citation_text
                or f"{citation.title}，p.{citation.page or '—'}，{citation.chunk_id}",
                numbered=False,
            )

    stream = io.BytesIO()
    document.save(stream)
    return stream.getvalue()


def _configure_docx_styles(document: Any) -> None:
    styles = document.styles
    normal = styles["Normal"]
    normal.font.name = _DOCX_FONT
    normal.font.size = Pt(11)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), _DOCX_FONT)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.1
    for name, size, color, before, after in (
        ("Heading 1", 16, _BLUE, 16, 8),
        ("Heading 2", 13, _BLUE, 12, 6),
        ("Heading 3", 12, _DARK_BLUE, 8, 4),
    ):
        style = styles[name]
        style.font.name = _DOCX_FONT
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color
        style._element.rPr.rFonts.set(qn("w:eastAsia"), _DOCX_FONT)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True


def _set_docx_header_footer(section: Any, report: ScientificReportPreviewDTO) -> None:
    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    _add_run(header, "NanoLoop - Scientific Analysis", 8.5, color=_MUTED)
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    _add_run(footer, f"{report.report_id}  |  ", 8.5, color=_MUTED)
    _add_field(footer, "PAGE")
    _add_run(footer, " / ", 8.5, color=_MUTED)
    _add_field(footer, "NUMPAGES")


def _add_heading(document: Any, text: str, *, level: int) -> None:
    document.add_heading(text, level=level)


def _add_callout(document: Any, text: str) -> None:
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(2)
    paragraph.paragraph_format.space_after = Pt(10)
    paragraph.paragraph_format.left_indent = Inches(0.14)
    paragraph.paragraph_format.right_indent = Inches(0.14)
    paragraph.paragraph_format.line_spacing = 1.15
    _paragraph_shading(paragraph, _CALLOUT_FILL)
    _paragraph_left_border(paragraph, "2E74B5", size=16, space=90)
    for index, line in enumerate(text.splitlines()):
        if index:
            paragraph.add_run().add_break()
        _add_run(paragraph, line, 10.5, color=_INK)


def _add_summary_with_overlay(document: Any, text: str, overlay: Path) -> None:
    table = document.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    left, right = table.rows[0].cells
    for cell, width in ((left, 4680), (right, 4680)):
        tc_pr = cell._tc.get_or_add_tcPr()
        tc_w = tc_pr.find(qn("w:tcW"))
        if tc_w is None:
            tc_w = OxmlElement("w:tcW")
            tc_pr.append(tc_w)
        tc_w.set(qn("w:type"), "dxa")
        tc_w.set(qn("w:w"), str(width))
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
        _set_cell_margins(cell, top=130, bottom=130, start=160, end=160)

    _set_cell_shading(left, _CALLOUT_FILL)
    summary = left.paragraphs[0]
    summary.paragraph_format.space_after = Pt(0)
    summary.paragraph_format.line_spacing = 1.12
    _paragraph_left_border(summary, "2E74B5", size=16, space=80)
    for index, line in enumerate(text.splitlines()):
        if index:
            summary.add_run().add_break()
        _add_run(summary, line, 9.5, color=_INK)

    figure_title = right.paragraphs[0]
    figure_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    figure_title.paragraph_format.space_after = Pt(4)
    _add_run(figure_title, "主运行识别叠加图", 9.5, color=_DARK_BLUE, bold=True)
    picture = right.add_paragraph()
    picture.alignment = WD_ALIGN_PARAGRAPH.CENTER
    picture.paragraph_format.space_after = Pt(3)
    picture.add_run().add_picture(str(overlay), width=Inches(2.95))
    caption = right.add_paragraph()
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption.paragraph_format.space_after = Pt(0)
    _add_run(caption, "图 1  模型覆盖与原图叠加，供人工核对。", 8, color=_MUTED)
    after = document.add_paragraph()
    after.paragraph_format.space_after = Pt(2)


def _add_metric_table(document: Any, report: ScientificReportPreviewDTO) -> None:
    headers = ("视野 / 模型", "质量", "颗粒数", "平均粒径", "覆盖率", "数密度")
    widths = (1800, 900, 900, 1700, 1100, 2960)
    table = document.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    for cell, header in zip(table.rows[0].cells, headers, strict=True):
        cell.text = header
        _style_table_cell(cell, header=True)
    for run in report.run_summaries:
        row = table.add_row()
        values = (
            (
                f"{run.filename or run.image_id}\n{run.model_id}"
                if report.analysis_mode == "batch"
                else run.model_id
            ),
            run.quality_status.value,
            str(run.particle_count),
            run.mean_equivalent_diameter,
            run.coverage,
            run.number_density,
        )
        for cell, value in zip(row.cells, values, strict=True):
            cell.text = value
            _style_table_cell(cell, header=False)
    _set_table_geometry(table, widths)
    after = document.add_paragraph(
        "表 1  所选运行的权威汇总；粒径与密度优先显示物理单位，缺少尺度时保留像素单位。"
    )
    after.paragraph_format.space_before = Pt(4)
    after.paragraph_format.space_after = Pt(8)
    for run in after.runs:
        _set_run_font(run, size=9, color=_MUTED)


def _add_batch_distribution_table(
    document: Any,
    batch: BatchResultSummaryDTO,
) -> None:
    headers = ("指标", "n", "均值", "标准差", "中位数", "IQR", "范围", "CV")
    widths = (1600, 650, 1000, 1000, 1000, 1300, 1550, 1260)
    table = document.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    for cell, header in zip(table.rows[0].cells, headers, strict=True):
        cell.text = header
        _style_table_cell(cell, header=True)
    for item in batch.distributions:
        row = table.add_row()
        values = (
            f"{item.label} ({item.unit})",
            str(item.sample_count),
            _number(item.mean),
            _number(item.std_dev),
            _number(item.median),
            f"{_number(item.q1)}–{_number(item.q3)}",
            f"{_number(item.minimum)}–{_number(item.maximum)}",
            (
                f"{_number(item.coefficient_of_variation * 100)}%"
                if item.coefficient_of_variation is not None
                else "—"
            ),
        )
        for cell, value in zip(row.cells, values, strict=True):
            cell.text = value
            _style_table_cell(cell, header=False)
    _set_table_geometry(table, widths)
    caption = document.add_paragraph(
        "表 2  运行级批量分布。异常筛查采用 1.5×IQR 规则，"
        f"共标记 {len(batch.outliers)} 个指标-运行组合。"
    )
    caption.paragraph_format.space_before = Pt(4)
    caption.paragraph_format.space_after = Pt(8)
    for run in caption.runs:
        _set_run_font(run, size=9, color=_MUTED)


def _add_provenance_table(document: Any, report: ScientificReportPreviewDTO) -> None:
    headers = ("运行", "模型版本", "尺度", "ROI", "threshold", "min area")
    widths = (1900, 1700, 1300, 1300, 1400, 1760)
    table = document.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    for cell, header in zip(table.rows[0].cells, headers, strict=True):
        cell.text = header
        _style_table_cell(cell, header=True)
    for item in report.provenance:
        row = table.add_row()
        values = (
            item.run_id,
            f"{item.model_id} / {item.model_version}",
            (
                f"{_number(item.scale_nm_per_pixel, digits=4)} nm/px"
                if item.scale_nm_per_pixel is not None
                else "pixel only"
            ),
            f"{item.roi_mode} / rev {item.box_revision if item.box_revision is not None else '—'}",
            _number(item.threshold, digits=4),
            f"{item.min_area_px} px",
        )
        for cell, value in zip(row.cells, values, strict=True):
            cell.text = value
            _style_table_cell(cell, header=False)
    _set_table_geometry(table, widths)


def _style_table_cell(cell: Any, *, header: bool) -> None:
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    _set_cell_shading(cell, _TABLE_HEADER_FILL if header else "FFFFFF")
    for paragraph in cell.paragraphs:
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.paragraph_format.line_spacing = 1.05
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if header else WD_ALIGN_PARAGRAPH.LEFT
        for run in paragraph.runs:
            _set_run_font(
                run,
                size=8.5 if not header else 9,
                color=_INK,
                bold=header,
            )


def _add_list_item(document: Any, text: str, *, numbered: bool) -> Any:
    paragraph = document.add_paragraph(style="List Number" if numbered else "List Bullet")
    paragraph.paragraph_format.left_indent = Inches(0.5)
    paragraph.paragraph_format.first_line_indent = Inches(-0.25)
    paragraph.paragraph_format.space_after = Pt(8 if numbered else 6)
    paragraph.paragraph_format.line_spacing = 1.167
    _add_run(paragraph, text, 10.5, color=_INK)
    return paragraph


def _set_table_geometry(table: Any, widths: Sequence[int]) -> None:
    if sum(widths) != 9360:
        raise ValueError("report table widths must total 9360 DXA")
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:type"), "dxa")
    tbl_w.set(qn("w:w"), "9360")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:type"), "dxa")
    tbl_ind.set(qn("w:w"), "120")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        column = OxmlElement("w:gridCol")
        column.set(qn("w:w"), str(width))
        grid.append(column)
    for row in table.rows:
        for cell, width in zip(row.cells, widths, strict=True):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:type"), "dxa")
            tc_w.set(qn("w:w"), str(width))
            _set_cell_margins(cell, top=80, bottom=80, start=120, end=120)
    header_props = table.rows[0]._tr.get_or_add_trPr()
    header = OxmlElement("w:tblHeader")
    header.set(qn("w:val"), "true")
    header_props.append(header)


def _set_cell_margins(cell: Any, **margins: int) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.find(qn("w:tcMar"))
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for key, value in margins.items():
        element = tc_mar.find(qn(f"w:{key}"))
        if element is None:
            element = OxmlElement(f"w:{key}")
            tc_mar.append(element)
        element.set(qn("w:w"), str(value))
        element.set(qn("w:type"), "dxa")


def _set_cell_shading(cell: Any, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = tc_pr.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        tc_pr.append(shading)
    shading.set(qn("w:fill"), fill)


def _paragraph_shading(paragraph: Any, fill: str) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    p_pr.append(shading)


def _paragraph_left_border(
    paragraph: Any,
    color: str,
    *,
    size: int,
    space: int,
) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    borders = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), str(size))
    left.set(qn("w:space"), str(space))
    left.set(qn("w:color"), color)
    borders.append(left)
    p_pr.append(borders)


def _add_run(
    paragraph: Any,
    text: str,
    size: float,
    *,
    color: RGBColor,
    bold: bool = False,
) -> Any:
    run = paragraph.add_run(text)
    _set_run_font(run, size=size, color=color, bold=bold)
    return run


def _set_run_font(
    run: Any,
    *,
    size: float,
    color: RGBColor,
    bold: bool = False,
) -> None:
    run.font.name = _DOCX_FONT
    run.font.size = Pt(size)
    run.font.color.rgb = color
    run.bold = bold
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), _DOCX_FONT)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), _DOCX_FONT)
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), _DOCX_FONT)


def _add_field(paragraph: Any, instruction: str) -> None:
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    text = OxmlElement("w:instrText")
    text.set(qn("xml:space"), "preserve")
    text.text = instruction
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    value = OxmlElement("w:t")
    value.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    for element in (begin, text, separate, value, end):
        run._r.append(element)
    _set_run_font(run, size=8.5, color=_MUTED)


def _render_pdf(report: ScientificReportPreviewDTO, *, overlay: Path | None) -> bytes:
    _ensure_pdf_font()
    stream = io.BytesIO()
    document = SimpleDocTemplate(
        stream,
        pagesize=letter,
        rightMargin=inch,
        leftMargin=inch,
        topMargin=0.75 * inch,
        bottomMargin=0.78 * inch,
        title=report.title,
        author="NanoLoop Agent",
        subject="SEM nanoparticle scientific analysis report",
    )
    styles = _pdf_styles()
    story: list[object] = [
        Paragraph("NANOLOOP SCIENTIFIC REPORT", styles["Kicker"]),
        Paragraph(_pdf_text(report.title), styles["Title"]),
        Paragraph("可追溯 SEM 分割、形貌统计与质量复核技术报告", styles["Subtitle"]),
        Spacer(1, 8),
    ]
    for label, value in (
        ("报告 ID", report.report_id),
        ("生成时间", report.generated_at.strftime("%Y-%m-%d %H:%M:%S UTC")),
        ("运行范围", "、".join(report.selected_run_ids)),
        (
            "综合方式",
            (
                f"本地模型 {report.synthesis_model}"
                if report.synthesis_provider == "local_llm"
                else "确定性可信模板"
            ),
        ),
    ):
        story.append(
            Paragraph(
                f"<b>{escape(label)}:</b> {escape(value)}",
                styles["Meta"],
            )
        )
    story.append(Paragraph("技术摘要", styles["H1"]))
    if overlay is not None:
        story.append(_pdf_summary_with_overlay(report, overlay, styles))
    else:
        story.append(_pdf_summary_callout(report, styles))
    story.append(Paragraph("主要结果与单位定义", styles["H1"]))
    for finding in report.findings:
        story.append(Paragraph(_pdf_text(finding.title), styles["H2"]))
        story.append(Paragraph(_pdf_text(finding.interpretation), styles["Body"]))
    story.extend(_pdf_metric_table(report, styles))
    if report.batch_summary is not None:
        story.append(Paragraph("批量分布与异常视野", styles["H2"]))
        story.extend(_pdf_batch_distribution_table(report.batch_summary, styles))
    story.append(Paragraph("范围、数据与方法", styles["H1"]))
    story.extend(_pdf_bullets(report.methodology, styles))
    story.append(Paragraph("限制、不确定性与稳健性", styles["H1"]))
    story.extend(_pdf_bullets(report.limitations, styles))
    story.append(Paragraph("建议的下一步", styles["H1"]))
    for recommendation in report.recommendations:
        story.append(
            KeepTogether(
                [
                    Paragraph(
                        f"<b>{recommendation.priority}. {_pdf_text(recommendation.action)}</b>",
                        styles["Body"],
                    ),
                    Paragraph(
                        f"<font color='#5C6474'>依据:</font> "
                        f"{_pdf_text(recommendation.rationale)}<br/>"
                        f"<font color='#5C6474'>验收:</font> "
                        f"{_pdf_text(recommendation.verification)}",
                        styles["Indented"],
                    ),
                ]
            )
        )
    story.append(Paragraph("仍需回答的问题", styles["H1"]))
    story.extend(_pdf_bullets(report.further_questions, styles))
    story.append(Paragraph("证据与可追溯性", styles["H1"]))
    story.extend(_pdf_provenance_table(report, styles))
    if report.data_evidence:
        story.append(Paragraph("确定性工具证据", styles["H2"]))
        for index, evidence_item in enumerate(report.data_evidence, start=1):
            source_ids = "、".join(evidence_item.source_run_ids) or "无来源运行"
            story.append(
                Paragraph(
                    f"<b>[D{index}] {_pdf_text(evidence_item.tool_name)}</b> - "
                    f"{_pdf_text(source_ids)}",
                    styles["Body"],
                )
            )
    if report.knowledge_citations:
        story.append(Paragraph("受管知识引用", styles["H2"]))
        story.extend(
            _pdf_bullets(
                [
                    citation.citation_text
                    or f"{citation.title}，p.{citation.page or '—'}，{citation.chunk_id}"
                    for citation in report.knowledge_citations
                ],
                styles,
            )
        )
    document.build(
        story,
        onFirstPage=lambda canvas, doc: _pdf_footer(canvas, doc, report),
        onLaterPages=lambda canvas, doc: _pdf_footer(canvas, doc, report),
    )
    return stream.getvalue()


def _ensure_pdf_font() -> None:
    with _PDF_FONT_LOCK:
        if _PDF_FONT not in pdfmetrics.getRegisteredFontNames():
            pdfmetrics.registerFont(UnicodeCIDFont(_PDF_FONT))


def _pdf_styles() -> dict[str, ParagraphStyle]:
    base = getSampleStyleSheet()
    return {
        "Kicker": ParagraphStyle(
            "ReportKicker",
            parent=base["Normal"],
            fontName=_PDF_FONT,
            fontSize=8,
            leading=10,
            textColor=colors.HexColor("#2E74B5"),
            spaceAfter=4,
        ),
        "Title": ParagraphStyle(
            "ReportTitle",
            parent=base["Title"],
            fontName=_PDF_FONT,
            fontSize=22,
            leading=28,
            alignment=TA_LEFT,
            textColor=colors.HexColor("#1D2330"),
            spaceAfter=5,
        ),
        "Subtitle": ParagraphStyle(
            "ReportSubtitle",
            parent=base["Normal"],
            fontName=_PDF_FONT,
            fontSize=11,
            leading=14,
            textColor=colors.HexColor("#5C6474"),
            spaceAfter=10,
        ),
        "Meta": ParagraphStyle(
            "ReportMeta",
            parent=base["Normal"],
            fontName=_PDF_FONT,
            fontSize=9,
            leading=12,
            textColor=colors.HexColor("#1D2330"),
            spaceAfter=2,
        ),
        "H1": ParagraphStyle(
            "ReportH1",
            parent=base["Heading1"],
            fontName=_PDF_FONT,
            fontSize=15,
            leading=19,
            textColor=colors.HexColor("#2E74B5"),
            spaceBefore=14,
            spaceAfter=7,
            keepWithNext=True,
        ),
        "H2": ParagraphStyle(
            "ReportH2",
            parent=base["Heading2"],
            fontName=_PDF_FONT,
            fontSize=11.5,
            leading=15,
            textColor=colors.HexColor("#1F4D78"),
            spaceBefore=8,
            spaceAfter=4,
            keepWithNext=True,
        ),
        "Body": ParagraphStyle(
            "ReportBody",
            parent=base["BodyText"],
            fontName=_PDF_FONT,
            fontSize=9.5,
            leading=13.5,
            textColor=colors.HexColor("#1D2330"),
            spaceAfter=6,
        ),
        "Indented": ParagraphStyle(
            "ReportIndented",
            parent=base["BodyText"],
            fontName=_PDF_FONT,
            fontSize=9,
            leading=13,
            leftIndent=18,
            textColor=colors.HexColor("#1D2330"),
            spaceAfter=7,
        ),
        "Callout": ParagraphStyle(
            "ReportCallout",
            parent=base["BodyText"],
            fontName=_PDF_FONT,
            fontSize=10,
            leading=15,
            textColor=colors.HexColor("#1D2330"),
        ),
        "Caption": ParagraphStyle(
            "ReportCaption",
            parent=base["Normal"],
            fontName=_PDF_FONT,
            fontSize=8,
            leading=10,
            alignment=TA_CENTER,
            textColor=colors.HexColor("#5C6474"),
            spaceBefore=4,
            spaceAfter=8,
        ),
        "FigureTitle": ParagraphStyle(
            "ReportFigureTitle",
            parent=base["Normal"],
            fontName=_PDF_FONT,
            fontSize=9,
            leading=11,
            alignment=TA_CENTER,
            textColor=colors.HexColor("#1F4D78"),
            spaceAfter=4,
        ),
        "FigureCaption": ParagraphStyle(
            "ReportFigureCaption",
            parent=base["Normal"],
            fontName=_PDF_FONT,
            fontSize=7,
            leading=9,
            alignment=TA_CENTER,
            textColor=colors.HexColor("#5C6474"),
            spaceBefore=3,
        ),
        "Table": ParagraphStyle(
            "ReportTable",
            parent=base["Normal"],
            fontName=_PDF_FONT,
            fontSize=7.3,
            leading=9,
            textColor=colors.HexColor("#1D2330"),
        ),
    }


def _pdf_summary_callout(
    report: ScientificReportPreviewDTO,
    styles: dict[str, ParagraphStyle],
) -> Table:
    return Table(
        [[Paragraph(_pdf_text(report.technical_summary), styles["Callout"])]],
        colWidths=[6.5 * inch],
        style=TableStyle(
            [
                ("BACKGROUND", (0, 0), (-1, -1), colors.HexColor("#F4F6F9")),
                ("BOX", (0, 0), (-1, -1), 0.5, colors.HexColor("#D8DEE8")),
                ("LINEBEFORE", (0, 0), (0, -1), 3, colors.HexColor("#2E74B5")),
                ("LEFTPADDING", (0, 0), (-1, -1), 12),
                ("RIGHTPADDING", (0, 0), (-1, -1), 12),
                ("TOPPADDING", (0, 0), (-1, -1), 10),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 10),
            ]
        ),
    )


def _pdf_summary_with_overlay(
    report: ScientificReportPreviewDTO,
    overlay: Path,
    styles: dict[str, ParagraphStyle],
) -> Table:
    summary = [Paragraph(_pdf_text(report.technical_summary), styles["Callout"])]
    figure = [
        Paragraph("主运行识别叠加图", styles["FigureTitle"]),
        _pdf_image(overlay, max_width=3.0 * inch, max_height=1.75 * inch),
        Paragraph("图 1  模型覆盖与原图叠加，供人工核对。", styles["FigureCaption"]),
    ]
    return Table(
        [[summary, figure]],
        colWidths=[3.25 * inch, 3.25 * inch],
        hAlign="LEFT",
        style=TableStyle(
            [
                ("BACKGROUND", (0, 0), (0, 0), colors.HexColor("#F4F6F9")),
                ("BOX", (0, 0), (-1, -1), 0.5, colors.HexColor("#D8DEE8")),
                ("LINEBEFORE", (0, 0), (0, 0), 3, colors.HexColor("#2E74B5")),
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
                ("LEFTPADDING", (0, 0), (-1, -1), 10),
                ("RIGHTPADDING", (0, 0), (-1, -1), 10),
                ("TOPPADDING", (0, 0), (-1, -1), 8),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 8),
            ]
        ),
    )


def _pdf_metric_table(
    report: ScientificReportPreviewDTO,
    styles: dict[str, ParagraphStyle],
) -> list[object]:
    rows: list[list[Paragraph]] = [
        [
            Paragraph("视野 / 模型", styles["Table"]),
            Paragraph("质量", styles["Table"]),
            Paragraph("颗粒数", styles["Table"]),
            Paragraph("平均粒径", styles["Table"]),
            Paragraph("覆盖率", styles["Table"]),
            Paragraph("数密度", styles["Table"]),
        ]
    ]
    for run in report.run_summaries:
        rows.append(
            [
                Paragraph(
                    _pdf_text(
                        f"{run.filename or run.image_id}\n{run.model_id}"
                        if report.analysis_mode == "batch"
                        else run.model_id
                    ),
                    styles["Table"],
                ),
                Paragraph(run.quality_status.value, styles["Table"]),
                Paragraph(str(run.particle_count), styles["Table"]),
                Paragraph(_pdf_text(run.mean_equivalent_diameter), styles["Table"]),
                Paragraph(run.coverage, styles["Table"]),
                Paragraph(_pdf_text(run.number_density), styles["Table"]),
            ]
        )
    table = Table(
        rows,
        colWidths=[1.25 * inch, 0.65 * inch, 0.6 * inch, 1.15 * inch, 0.7 * inch, 2.15 * inch],
        repeatRows=1,
        hAlign="LEFT",
    )
    table.setStyle(
        TableStyle(
            [
                ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#F2F4F7")),
                ("GRID", (0, 0), (-1, -1), 0.35, colors.HexColor("#D9DEE8")),
                ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
                ("LEFTPADDING", (0, 0), (-1, -1), 5),
                ("RIGHTPADDING", (0, 0), (-1, -1), 5),
                ("TOPPADDING", (0, 0), (-1, -1), 6),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 6),
            ]
        )
    )
    return [
        table,
        Paragraph(
            "表 1  所选运行的权威汇总；缺少尺度时保留像素单位。",
            styles["Caption"],
        ),
    ]


def _pdf_batch_distribution_table(
    batch: BatchResultSummaryDTO,
    styles: dict[str, ParagraphStyle],
) -> list[object]:
    rows: list[list[Paragraph]] = [
        [
            Paragraph("指标", styles["Table"]),
            Paragraph("n", styles["Table"]),
            Paragraph("均值", styles["Table"]),
            Paragraph("标准差", styles["Table"]),
            Paragraph("中位数", styles["Table"]),
            Paragraph("IQR", styles["Table"]),
            Paragraph("范围", styles["Table"]),
            Paragraph("CV", styles["Table"]),
        ]
    ]
    for item in batch.distributions:
        rows.append(
            [
                Paragraph(_pdf_text(f"{item.label} ({item.unit})"), styles["Table"]),
                Paragraph(str(item.sample_count), styles["Table"]),
                Paragraph(_number(item.mean), styles["Table"]),
                Paragraph(_number(item.std_dev), styles["Table"]),
                Paragraph(_number(item.median), styles["Table"]),
                Paragraph(f"{_number(item.q1)}–{_number(item.q3)}", styles["Table"]),
                Paragraph(
                    f"{_number(item.minimum)}–{_number(item.maximum)}",
                    styles["Table"],
                ),
                Paragraph(
                    (
                        f"{_number(item.coefficient_of_variation * 100)}%"
                        if item.coefficient_of_variation is not None
                        else "—"
                    ),
                    styles["Table"],
                ),
            ]
        )
    table = Table(
        rows,
        colWidths=[
            1.1 * inch,
            0.35 * inch,
            0.7 * inch,
            0.7 * inch,
            0.7 * inch,
            0.9 * inch,
            1.1 * inch,
            0.95 * inch,
        ],
        repeatRows=1,
        hAlign="LEFT",
    )
    table.setStyle(
        TableStyle(
            [
                ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#F2F4F7")),
                ("GRID", (0, 0), (-1, -1), 0.35, colors.HexColor("#D9DEE8")),
                ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
                ("LEFTPADDING", (0, 0), (-1, -1), 4),
                ("RIGHTPADDING", (0, 0), (-1, -1), 4),
                ("TOPPADDING", (0, 0), (-1, -1), 5),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 5),
            ]
        )
    )
    return [
        table,
        Paragraph(
            f"表 2  运行级分布；1.5×IQR 筛查标记 {len(batch.outliers)} 个指标-运行组合。",
            styles["Caption"],
        ),
    ]


def _pdf_provenance_table(
    report: ScientificReportPreviewDTO,
    styles: dict[str, ParagraphStyle],
) -> list[object]:
    rows: list[list[Paragraph]] = [
        [
            Paragraph("运行", styles["Table"]),
            Paragraph("模型版本", styles["Table"]),
            Paragraph("尺度", styles["Table"]),
            Paragraph("ROI", styles["Table"]),
            Paragraph("threshold", styles["Table"]),
            Paragraph("min area", styles["Table"]),
        ]
    ]
    for item in report.provenance:
        rows.append(
            [
                Paragraph(escape(item.run_id), styles["Table"]),
                Paragraph(
                    _pdf_text(f"{item.model_id} / {item.model_version}"),
                    styles["Table"],
                ),
                Paragraph(
                    (
                        f"{_number(item.scale_nm_per_pixel, digits=4)} nm/px"
                        if item.scale_nm_per_pixel is not None
                        else "pixel only"
                    ),
                    styles["Table"],
                ),
                Paragraph(
                    _pdf_text(
                        f"{item.roi_mode} / rev "
                        f"{item.box_revision if item.box_revision is not None else '—'}"
                    ),
                    styles["Table"],
                ),
                Paragraph(_number(item.threshold, digits=4), styles["Table"]),
                Paragraph(f"{item.min_area_px} px", styles["Table"]),
            ]
        )
    table = Table(
        rows,
        colWidths=[1.2 * inch, 1.7 * inch, 1.0 * inch, 1.1 * inch, 0.75 * inch, 0.75 * inch],
        repeatRows=1,
        hAlign="LEFT",
    )
    table.setStyle(
        TableStyle(
            [
                ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#F2F4F7")),
                ("GRID", (0, 0), (-1, -1), 0.35, colors.HexColor("#D9DEE8")),
                ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
                ("LEFTPADDING", (0, 0), (-1, -1), 5),
                ("RIGHTPADDING", (0, 0), (-1, -1), 5),
                ("TOPPADDING", (0, 0), (-1, -1), 6),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 6),
            ]
        )
    )
    return [table]


def _pdf_bullets(
    values: Iterable[str],
    styles: dict[str, ParagraphStyle],
) -> list[Paragraph]:
    return [Paragraph(f"- {_pdf_text(value)}", styles["Body"]) for value in values]


def _pdf_image(
    path: Path,
    *,
    max_width: float = 6.35 * inch,
    max_height: float = 3.6 * inch,
) -> Any:
    image = Image(str(path))
    ratio = min(max_width / image.imageWidth, max_height / image.imageHeight)
    image.drawWidth = image.imageWidth * ratio
    image.drawHeight = image.imageHeight * ratio
    image.hAlign = "CENTER"
    return image


def _pdf_footer(canvas: Any, document: Any, report: ScientificReportPreviewDTO) -> None:
    canvas.saveState()
    canvas.setFont(_PDF_FONT, 7.5)
    canvas.setFillColor(colors.HexColor("#737B8C"))
    canvas.drawString(inch, 0.55 * inch, report.report_id)
    canvas.drawRightString(7.5 * inch, 0.55 * inch, f"第 {document.page} 页")
    canvas.restoreState()


def _pdf_text(value: str) -> str:
    portable = (
        value.replace("·", "-")
        .replace("µ", "u")
        .replace("μ", "u")
        .replace("⁻²", "^-2")
        .replace("⁻¹", "^-1")
        .replace("²", "^2")
        .replace("¹", "^1")
        .replace("⁻", "^-")
    )
    return escape(portable).replace("\n", "<br/>")


def _human_reason(reason: str) -> str:
    labels = {
        "small_fragment_ratio_high": "小碎片比例偏高",
        "foreground_ratio_low": "前景覆盖率偏低",
        "foreground_ratio_high": "前景覆盖率偏高",
        "border_touch_ratio_high": "边界接触比例偏高",
        "agglomeration_ratio_high": "粘连/团聚比例偏高",
    }
    return labels.get(reason, reason.replace("_", " "))


def _number(value: float | int | None, *, digits: int = 2) -> str:
    if value is None:
        return "—"
    if digits == 0:
        return f"{value:,.0f}"
    return f"{value:,.{digits}f}".rstrip("0").rstrip(".")


__all__ = ["ScientificReportBuilder"]
