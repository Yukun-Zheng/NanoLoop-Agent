#!/usr/bin/env python3
"""Build polished A4 DOCX sources from the submission Markdown files."""

from __future__ import annotations

import re
import sys
from collections.abc import Iterable
from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[2]
GENERATED = ROOT / "submission" / "generated"
GENERATED.mkdir(parents=True, exist_ok=True)

FONT_CJK = "Hiragino Sans GB"
FONT_LATIN = "Hiragino Sans GB"
FONT_MONO = "Menlo"

INK = "17212B"
MUTED = "5F6B76"
BLUE = "2E74B5"
BLUE_DARK = "1F4D78"
BLUE_PALE = "E8EEF5"
GREEN = "356859"
GREEN_PALE = "EDF5F1"
AMBER = "7A5A00"
AMBER_PALE = "F8F2E5"
LINE = "D5DCE3"
LIGHT = "F4F6F9"
WHITE = "FFFFFF"


def read_source(path: Path) -> tuple[dict[str, str], list[str]]:
    lines = path.read_text(encoding="utf-8").splitlines()
    metadata: dict[str, str] = {}
    if lines and lines[0].strip() == "---":
        end = lines.index("---", 1)
        for line in lines[1:end]:
            if ":" not in line:
                continue
            key, value = line.split(":", 1)
            metadata[key.strip()] = value.strip()
        lines = lines[end + 1 :]
    return metadata, lines


def set_font(
    run,
    size: float | None = None,
    bold: bool | None = None,
    color: str | None = None,
    mono: bool = False,
) -> None:
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    font_name = FONT_MONO if mono else FONT_LATIN
    cjk_name = FONT_MONO if mono else FONT_CJK
    run.font.name = font_name
    r_fonts = run._element.get_or_add_rPr().rFonts
    r_fonts.set(qn("w:ascii"), font_name)
    r_fonts.set(qn("w:hAnsi"), font_name)
    r_fonts.set(qn("w:cs"), font_name)
    r_fonts.set(qn("w:eastAsia"), cjk_name)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_border(cell, **edges: dict[str, str]) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = tc_pr.first_child_found_in("w:tcBorders")
    if tc_borders is None:
        tc_borders = OxmlElement("w:tcBorders")
        tc_pr.append(tc_borders)
    for edge_name, edge_data in edges.items():
        tag = f"w:{edge_name}"
        edge = tc_borders.find(qn(tag))
        if edge is None:
            edge = OxmlElement(tag)
            tc_borders.append(edge)
        for key in ("val", "sz", "space", "color"):
            if key in edge_data:
                edge.set(qn(f"w:{key}"), str(edge_data[key]))


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def prevent_row_split(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    tr_pr.append(cant_split)


def set_table_layout(table, widths: list[int]) -> None:
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")
    tbl_width = tbl_pr.find(qn("w:tblW"))
    if tbl_width is None:
        tbl_width = OxmlElement("w:tblW")
        tbl_pr.append(tbl_width)
    tbl_width.set(qn("w:type"), "dxa")
    tbl_width.set(qn("w:w"), str(sum(widths)))
    tbl_indent = tbl_pr.find(qn("w:tblInd"))
    if tbl_indent is None:
        tbl_indent = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_indent)
    tbl_indent.set(qn("w:type"), "dxa")
    tbl_indent.set(qn("w:w"), "120")
    cell_margins = tbl_pr.find(qn("w:tblCellMar"))
    if cell_margins is None:
        cell_margins = OxmlElement("w:tblCellMar")
        tbl_pr.append(cell_margins)
    for edge_name, value in (("top", 80), ("bottom", 80), ("start", 120), ("end", 120)):
        edge = cell_margins.find(qn(f"w:{edge_name}"))
        if edge is None:
            edge = OxmlElement(f"w:{edge_name}")
            cell_margins.append(edge)
        edge.set(qn("w:type"), "dxa")
        edge.set(qn("w:w"), str(value))
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for cell, width in zip(row.cells, widths, strict=True):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:type"), "dxa")
            tc_w.set(qn("w:w"), str(width))


def set_paragraph_keep(paragraph, *, next_: bool = False, lines: bool = False) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    if next_:
        keep_next = OxmlElement("w:keepNext")
        p_pr.append(keep_next)
    if lines:
        keep_lines = OxmlElement("w:keepLines")
        p_pr.append(keep_lines)


def set_outline_level(paragraph, level: int) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    outline = OxmlElement("w:outlineLvl")
    outline.set(qn("w:val"), str(level))
    p_pr.append(outline)


def add_field(run, instruction: str) -> None:
    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = instruction
    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char_begin)
    run._r.append(instr_text)
    run._r.append(fld_char_end)


def add_numbering_definition(document: Document, *, fmt: str, text: str) -> int:
    numbering = document.part.numbering_part.element
    abstract_ids = [
        int(element.get(qn("w:abstractNumId")))
        for element in numbering.findall(qn("w:abstractNum"))
    ]
    num_ids = [int(element.get(qn("w:numId"))) for element in numbering.findall(qn("w:num"))]
    abstract_id = max(abstract_ids, default=0) + 1
    num_id = max(num_ids, default=0) + 1

    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)
    level = OxmlElement("w:lvl")
    level.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    level.append(start)
    num_fmt = OxmlElement("w:numFmt")
    num_fmt.set(qn("w:val"), fmt)
    level.append(num_fmt)
    level_text = OxmlElement("w:lvlText")
    level_text.set(qn("w:val"), text)
    level.append(level_text)
    justification = OxmlElement("w:lvlJc")
    justification.set(qn("w:val"), "left")
    level.append(justification)
    p_pr = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), "540")
    tabs.append(tab)
    p_pr.append(tabs)
    indent = OxmlElement("w:ind")
    indent.set(qn("w:left"), "540")
    indent.set(qn("w:hanging"), "279")
    p_pr.append(indent)
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:after"), "80")
    spacing.set(qn("w:line"), "290")
    spacing.set(qn("w:lineRule"), "auto")
    p_pr.append(spacing)
    level.append(p_pr)
    abstract.append(level)
    numbering.append(abstract)

    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(abstract_id))
    num.append(abstract_ref)
    numbering.append(num)
    return num_id


def apply_numbering(paragraph, num_id: int) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = OxmlElement("w:numPr")
    level = OxmlElement("w:ilvl")
    level.set(qn("w:val"), "0")
    num = OxmlElement("w:numId")
    num.set(qn("w:val"), str(num_id))
    num_pr.append(level)
    num_pr.append(num)
    p_pr.append(num_pr)


def configure_document(document: Document, title: str) -> tuple[int, int]:
    section = document.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    section.different_first_page_header_footer = True

    styles = document.styles
    normal = styles["Normal"]
    normal.font.name = FONT_LATIN
    for attribute, value in [
        ("w:ascii", FONT_LATIN),
        ("w:hAnsi", FONT_LATIN),
        ("w:cs", FONT_LATIN),
        ("w:eastAsia", FONT_CJK),
    ]:
        normal._element.rPr.rFonts.set(qn(attribute), value)
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(8)
    normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    normal.paragraph_format.line_spacing = 1.333

    for style_name, size, color, before, after in [
        ("Title", 30, INK, 0, 12),
        ("Subtitle", 15, MUTED, 0, 10),
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, BLUE_DARK, 8, 4),
        ("Caption", 9, MUTED, 4, 8),
    ]:
        style = styles[style_name]
        style.font.name = FONT_LATIN
        for attribute, value in [
            ("w:ascii", FONT_LATIN),
            ("w:hAnsi", FONT_LATIN),
            ("w:cs", FONT_LATIN),
            ("w:eastAsia", FONT_CJK),
        ]:
            style._element.rPr.rFonts.set(qn(attribute), value)
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        if style_name.startswith("Heading"):
            style.font.bold = True
            style.paragraph_format.keep_with_next = True
            style.paragraph_format.keep_together = True

    document.core_properties.title = title
    document.core_properties.subject = "2026 年首届深圳大学 AI4S 智能体创新大赛参赛材料"
    document.core_properties.author = ""
    document.core_properties.last_modified_by = ""
    document.core_properties.comments = ""

    header = section.header
    header.is_linked_to_previous = False
    p = header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run(title)
    set_font(run, size=8.2, color=MUTED)
    p.paragraph_format.space_after = Pt(0)
    p_pr = p._p.get_or_add_pPr()
    border = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "5")
    bottom.set(qn("w:color"), LINE)
    border.append(bottom)
    p_pr.append(border)

    footer = section.footer
    footer.is_linked_to_previous = False
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("NanoLoop · 智能体使用手册   ")
    set_font(run, size=8.1, color=MUTED)
    page_run = p.add_run()
    add_field(page_run, "PAGE")
    set_font(page_run, size=8.1, color=MUTED)
    bullet_num_id = add_numbering_definition(document, fmt="bullet", text="•")
    decimal_num_id = add_numbering_definition(document, fmt="decimal", text="%1.")
    return bullet_num_id, decimal_num_id


def add_cover(document: Document, metadata: dict[str, str], source: Path) -> None:
    p = document.add_paragraph()
    p.paragraph_format.space_after = Pt(10)
    run = p.add_run(metadata.get("kicker", "AI4S"))
    set_font(run, size=10.5, bold=True, color=BLUE_DARK)

    accent = document.add_table(rows=1, cols=1)
    accent.alignment = WD_ALIGN_PARAGRAPH.LEFT
    cell = accent.cell(0, 0)
    set_cell_shading(cell, BLUE)
    cell.width = Cm(16.5)
    cell.height = Cm(0.12)
    cell.text = ""
    set_cell_border(
        cell, top={"val": "nil"}, bottom={"val": "nil"}, left={"val": "nil"}, right={"val": "nil"}
    )
    document.add_paragraph().paragraph_format.space_after = Pt(0)

    p = document.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(5)
    cover_title = metadata.get("title", "NanoLoop")
    run = p.add_run(cover_title)
    # Long Chinese titles otherwise leave a single orphan character on the next
    # line in LibreOffice's PDF renderer. Keep the title prominent, but fit it
    # as a deliberate one- or two-line lockup.
    cover_title_size = 28 if len(cover_title) >= 20 else 33
    set_font(run, size=cover_title_size, bold=True, color=INK)

    p = document.add_paragraph()
    p.paragraph_format.space_after = Pt(9)
    run = p.add_run(metadata.get("subtitle", ""))
    set_font(run, size=15, color=MUTED)

    badge = document.add_table(rows=1, cols=1)
    badge.autofit = False
    badge.alignment = WD_ALIGN_PARAGRAPH.LEFT
    cell = badge.cell(0, 0)
    set_cell_shading(cell, BLUE_PALE)
    set_cell_border(
        cell,
        top={"val": "single", "sz": "8", "color": "C9D1FF"},
        bottom={"val": "single", "sz": "8", "color": "C9D1FF"},
        left={"val": "single", "sz": "8", "color": "C9D1FF"},
        right={"val": "single", "sz": "8", "color": "C9D1FF"},
    )
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(metadata.get("badge", ""))
    set_font(run, size=10, bold=True, color=BLUE_DARK)

    hero_text = metadata.get("hero")
    if hero_text:
        hero = (source.parent / hero_text).resolve()
        if hero.exists():
            p = document.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(10)
            p.paragraph_format.space_after = Pt(8)
            hero_shape = p.add_run().add_picture(str(hero), width=Cm(16.2))
            hero_shape._inline.docPr.set("descr", "NanoLoop 结果复核界面")
            hero_shape._inline.docPr.set("title", "NanoLoop 结果复核界面")

    info = document.add_table(rows=2, cols=2)
    info.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_table_layout(info, [4500, 4860])
    values = [
        (metadata.get("team", ""), metadata.get("badge", "")),
        (metadata.get("date", ""), "单一手册 · 离线 Docker 交付"),
    ]
    for row, values_row in zip(info.rows, values, strict=True):
        for cell, value in zip(row.cells, values_row, strict=True):
            set_cell_shading(cell, LIGHT)
            set_cell_border(
                cell,
                top={"val": "single", "sz": "4", "color": LINE},
                bottom={"val": "single", "sz": "4", "color": LINE},
                left={"val": "single", "sz": "4", "color": LINE},
                right={"val": "single", "sz": "4", "color": LINE},
            )
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            run = p.add_run(value)
            set_font(run, size=9.2, bold=row is info.rows[0], color=INK)

    promise = document.add_table(rows=1, cols=1)
    promise.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_table_layout(promise, [9360])
    cell = promise.cell(0, 0)
    set_cell_shading(cell, INK)
    set_cell_border(
        cell,
        top={"val": "single", "sz": "4", "color": INK},
        bottom={"val": "single", "sz": "4", "color": INK},
        left={"val": "single", "sz": "4", "color": INK},
        right={"val": "single", "sz": "4", "color": INK},
    )
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(7)
    p.paragraph_format.space_after = Pt(7)
    run = p.add_run("原始输入可核对 · 科学配置可复现 · 质量边界不隐藏 · 结果能够回到实验")
    set_font(run, size=10.2, bold=True, color=WHITE)


def extract_headings(lines: list[str]) -> list[str]:
    return [line[2:].strip() for line in lines if line.startswith("# ")][:24]


def add_contents_overview(document: Document, headings: list[str]) -> None:
    p = document.add_paragraph()
    p.paragraph_format.page_break_before = True
    run = p.add_run("阅读路线")
    set_font(run, size=25, bold=True, color=INK)
    p.paragraph_format.space_after = Pt(4)
    p = document.add_paragraph()
    run = p.add_run("先读问题与闭环，再看技术边界、证据和部署。")
    set_font(run, size=10.5, color=MUTED)
    p.paragraph_format.space_after = Pt(12)

    cols = 2 if len(headings) > 12 else 1
    rows = (len(headings) + cols - 1) // cols
    table = document.add_table(rows=rows, cols=cols)
    table.alignment = WD_ALIGN_PARAGRAPH.CENTER
    widths = [4680] * cols if cols == 2 else [9360]
    set_table_layout(table, widths)
    for idx, heading_text in enumerate(headings):
        row_idx = idx % rows
        col_idx = idx // rows
        cell = table.cell(row_idx, col_idx)
        set_cell_shading(cell, WHITE if row_idx % 2 else LIGHT)
        set_cell_border(
            cell,
            bottom={"val": "single", "sz": "4", "color": LINE},
            top={"val": "nil"},
            left={"val": "nil"},
            right={"val": "nil"},
        )
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(f"{idx + 1:02d}")
        set_font(run, size=8.5, bold=True, color=BLUE)
        run = p.add_run(f"   {heading_text}")
        set_font(run, size=10.2, color=INK)

    p = document.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run("证据原则")
    set_font(run, size=9.5, bold=True, color=GREEN)
    p = document.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run("工程闭环与科学泛化分开表述；缺少尺度、模型或引用时，系统明确降级。")
    set_font(run, size=10.2, color=MUTED)

    scoring = document.add_table(rows=2, cols=4)
    scoring.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_table_layout(scoring, [2340, 2340, 2340, 2340])
    labels = [
        ("30%", "科学及应用价值"),
        ("30%", "技术深度"),
        ("20%", "技术落地性"),
        ("20%", "演示效果"),
    ]
    for col, (score, label) in enumerate(labels):
        top_cell = scoring.cell(0, col)
        bottom_cell = scoring.cell(1, col)
        set_cell_shading(top_cell, BLUE_PALE)
        set_cell_shading(bottom_cell, LIGHT)
        for cell in (top_cell, bottom_cell):
            set_cell_border(
                cell,
                top={"val": "single", "sz": "4", "color": LINE},
                bottom={"val": "single", "sz": "4", "color": LINE},
                left={"val": "single", "sz": "4", "color": LINE},
                right={"val": "single", "sz": "4", "color": LINE},
            )
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            cell.paragraphs[0].paragraph_format.space_after = Pt(0)
        run = top_cell.paragraphs[0].add_run(score)
        set_font(run, size=15, bold=True, color=BLUE_DARK)
        run = bottom_cell.paragraphs[0].add_run(label)
        set_font(run, size=8.8, bold=True, color=INK)
    document.add_page_break()


INLINE_PATTERN = re.compile(r"(\*\*.+?\*\*|`[^`]+`)")


def add_inline(
    paragraph, text: str, *, size: float | None = None, color: str | None = None
) -> None:
    for chunk in INLINE_PATTERN.split(text):
        if not chunk:
            continue
        if chunk.startswith("**") and chunk.endswith("**"):
            run = paragraph.add_run(chunk[2:-2])
            set_font(run, size=size, bold=True, color=color or INK)
        elif chunk.startswith("`") and chunk.endswith("`"):
            run = paragraph.add_run(chunk[1:-1])
            set_font(run, size=(size or 10) - 0.4, color=BLUE_DARK, mono=True)
            shading = OxmlElement("w:shd")
            shading.set(qn("w:fill"), BLUE_PALE)
            run._r.get_or_add_rPr().append(shading)
        else:
            run = paragraph.add_run(chunk)
            set_font(run, size=size, color=color or INK)


def add_callout(document: Document, text: str) -> None:
    table = document.add_table(rows=1, cols=1)
    table.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_table_layout(table, [9360])
    cell = table.cell(0, 0)
    set_cell_shading(cell, BLUE_PALE)
    set_cell_border(
        cell,
        left={"val": "single", "sz": "22", "color": BLUE},
        top={"val": "single", "sz": "4", "color": "D5DBFF"},
        right={"val": "single", "sz": "4", "color": "D5DBFF"},
        bottom={"val": "single", "sz": "4", "color": "D5DBFF"},
    )
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    add_inline(p, text, size=10.2, color=INK)


def parse_table_rows(lines: list[str]) -> list[list[str]]:
    rows: list[list[str]] = []
    for line in lines:
        stripped = line.strip().strip("|")
        rows.append([cell.strip() for cell in stripped.split("|")])
    return rows


def is_table_separator(line: str) -> bool:
    cells = [cell.strip() for cell in line.strip().strip("|").split("|")]
    return bool(cells) and all(re.fullmatch(r":?-{3,}:?", cell) for cell in cells)


def add_markdown_table(document: Document, lines: list[str]) -> None:
    rows = parse_table_rows(lines)
    if len(rows) >= 2 and is_table_separator(lines[1]):
        rows.pop(1)
    column_count = max(len(row) for row in rows)
    table = document.add_table(rows=len(rows), cols=column_count)
    table.alignment = WD_ALIGN_PARAGRAPH.CENTER
    table.style = "Table Grid"
    total_width = 9360
    if column_count == 2:
        widths = [2700, 6660]
    elif column_count == 3:
        widths = [2160, 3600, 3600]
    elif column_count == 4:
        widths = [1600, 2500, 2630, 2630]
    else:
        widths = [total_width // column_count] * column_count
        widths[-1] += total_width - sum(widths)
    set_table_layout(table, widths)
    for row_idx, (row, source_row) in enumerate(zip(table.rows, rows, strict=True)):
        prevent_row_split(row)
        if row_idx == 0:
            set_repeat_table_header(row)
        for col_idx, cell in enumerate(row.cells):
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            if row_idx == 0:
                set_cell_shading(cell, BLUE_DARK)
            elif row_idx % 2 == 0:
                set_cell_shading(cell, LIGHT)
            else:
                set_cell_shading(cell, WHITE)
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_after = Pt(1.5)
                paragraph.paragraph_format.space_before = Pt(1.5)
                text = source_row[col_idx] if col_idx < len(source_row) else ""
                paragraph.clear()
                add_inline(
                    paragraph,
                    text,
                    size=8.4 if column_count >= 4 else 9,
                    color=WHITE if row_idx == 0 else INK,
                )
                if row_idx == 0:
                    for run in paragraph.runs:
                        run.bold = True
            set_cell_border(
                cell,
                top={"val": "single", "sz": "4", "color": LINE},
                bottom={"val": "single", "sz": "4", "color": LINE},
                left={"val": "single", "sz": "4", "color": LINE},
                right={"val": "single", "sz": "4", "color": LINE},
            )


def add_code_block(document: Document, lines: list[str]) -> None:
    table = document.add_table(rows=1, cols=1)
    table.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_table_layout(table, [9360])
    cell = table.cell(0, 0)
    set_cell_shading(cell, "20232B")
    set_cell_border(
        cell,
        top={"val": "single", "sz": "4", "color": "333845"},
        bottom={"val": "single", "sz": "4", "color": "333845"},
        left={"val": "single", "sz": "4", "color": "333845"},
        right={"val": "single", "sz": "4", "color": "333845"},
    )
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.08
    for idx, line in enumerate(lines):
        run = p.add_run(line)
        set_font(run, size=8, color="F4F6FA", mono=True)
        if idx != len(lines) - 1:
            run.add_break()


def add_image(document: Document, source_path: Path, alt_text: str) -> None:
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(7)
    p.paragraph_format.space_after = Pt(2)
    picture = p.add_run().add_picture(str(source_path), width=Cm(16.2))
    picture._inline.docPr.set("descr", alt_text)
    picture._inline.docPr.set("title", alt_text)
    set_paragraph_keep(p, lines=True)
    caption = document.add_paragraph(style="Caption")
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_inline(caption, f"图｜{alt_text}", size=8.7, color=MUTED)
    set_paragraph_keep(caption, lines=True)


def render_content(
    document: Document,
    source: Path,
    lines: list[str],
    *,
    bullet_num_id: int,
    decimal_num_id: int,
) -> None:
    i = 0
    while i < len(lines):
        raw = lines[i]
        stripped = raw.strip()
        if not stripped:
            i += 1
            continue
        if stripped == "\\newpage":
            document.add_page_break()
            i += 1
            continue
        if stripped.startswith("```"):
            i += 1
            code_lines: list[str] = []
            while i < len(lines) and not lines[i].strip().startswith("```"):
                code_lines.append(lines[i])
                i += 1
            i += 1
            add_code_block(document, code_lines or [""])
            continue
        heading_match = re.match(r"^(#{1,3})\s+(.+)$", raw)
        if heading_match:
            level = len(heading_match.group(1))
            text = heading_match.group(2).strip()
            paragraph = document.add_paragraph(style=f"Heading {level}")
            paragraph.clear()
            add_inline(paragraph, text, color=BLUE_DARK if level == 1 else INK)
            set_outline_level(paragraph, level - 1)
            set_paragraph_keep(paragraph, next_=True, lines=True)
            i += 1
            continue
        image_match = re.fullmatch(r"!\[(.*?)\]\((.*?)\)", stripped)
        if image_match:
            image_path = (source.parent / image_match.group(2)).resolve()
            if not image_path.exists():
                raise FileNotFoundError(f"Missing image {image_path} referenced by {source}")
            add_image(document, image_path, image_match.group(1))
            i += 1
            continue
        if stripped.startswith(">"):
            quote_lines: list[str] = []
            while i < len(lines) and lines[i].strip().startswith(">"):
                quote_lines.append(lines[i].strip()[1:].strip())
                i += 1
            add_callout(document, " ".join(quote_lines))
            continue
        if "|" in raw and i + 1 < len(lines) and is_table_separator(lines[i + 1]):
            table_lines = [raw, lines[i + 1]]
            i += 2
            while i < len(lines) and "|" in lines[i] and lines[i].strip():
                table_lines.append(lines[i])
                i += 1
            add_markdown_table(document, table_lines)
            continue
        bullet_match = re.match(r"^-\s+(.*)$", stripped)
        if bullet_match:
            paragraph = document.add_paragraph()
            paragraph.paragraph_format.left_indent = Inches(0.375)
            paragraph.paragraph_format.first_line_indent = Inches(-0.194)
            paragraph.paragraph_format.space_after = Pt(4)
            paragraph.paragraph_format.line_spacing = 1.208
            apply_numbering(paragraph, bullet_num_id)
            add_inline(paragraph, bullet_match.group(1))
            set_paragraph_keep(paragraph, lines=True)
            i += 1
            continue
        number_match = re.match(r"^\d+\.\s+(.*)$", stripped)
        if number_match:
            paragraph = document.add_paragraph()
            paragraph.paragraph_format.left_indent = Inches(0.375)
            paragraph.paragraph_format.first_line_indent = Inches(-0.194)
            paragraph.paragraph_format.space_after = Pt(4)
            paragraph.paragraph_format.line_spacing = 1.208
            apply_numbering(paragraph, decimal_num_id)
            add_inline(paragraph, number_match.group(1))
            set_paragraph_keep(paragraph, lines=True)
            i += 1
            continue

        paragraph_lines = [stripped]
        i += 1
        while i < len(lines):
            nxt = lines[i].strip()
            if (
                not nxt
                or nxt == "\\newpage"
                or nxt.startswith("#")
                or nxt.startswith("```")
                or nxt.startswith(">")
                or re.fullmatch(r"!\[(.*?)\]\((.*?)\)", nxt)
                or re.match(r"^-\s+", nxt)
                or re.match(r"^\d+\.\s+", nxt)
                or ("|" in nxt and i + 1 < len(lines) and is_table_separator(lines[i + 1]))
            ):
                break
            paragraph_lines.append(nxt)
            i += 1
        paragraph = document.add_paragraph()
        add_inline(paragraph, " ".join(paragraph_lines))
        set_paragraph_keep(paragraph, lines=True)


def build(source: Path) -> Path:
    metadata, lines = read_source(source)
    document = Document()
    title = metadata.get("title", source.stem)
    bullet_num_id, decimal_num_id = configure_document(document, title)
    add_cover(document, metadata, source)
    add_contents_overview(document, extract_headings(lines))
    render_content(
        document,
        source,
        lines,
        bullet_num_id=bullet_num_id,
        decimal_num_id=decimal_num_id,
    )

    for paragraph in document.paragraphs:
        if paragraph.style.name.startswith("Heading"):
            set_paragraph_keep(paragraph, next_=True, lines=True)

    output = GENERATED / f"{source.stem}.docx"
    document.save(output)
    print(f"{output.relative_to(ROOT)}\t{output.stat().st_size} bytes")
    return output


def main(argv: Iterable[str]) -> int:
    sources = list(argv)
    if not sources:
        sources = [
            "submission/docs/NanoLoop智能体设计文档.md",
            "submission/docs/NanoLoop_Docker部署与使用手册.md",
            "submission/video/NanoLoop三分钟演示录制手册.md",
        ]
    for value in sources:
        build((ROOT / value).resolve())
    return 0


if __name__ == "__main__":
    raise SystemExit(main(sys.argv[1:]))
